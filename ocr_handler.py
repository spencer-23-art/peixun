import os
import re
import uuid
import shutil
import time
import datetime
import sqlite3
import cv2
import numpy as np
from PIL import Image, ImageOps
from docx import Document
from docx.shared import Inches
from modelscope import snapshot_download
from rapidocr_onnxruntime import RapidOCR
from app.postprocess import NATIONS
from app.config import DB_PATH, UPLOAD_DIR, TEMP_IDS_DIR, CARDS_DIR


# ================= 1. PP-OCRv6 Tiny 模型加载与 OCR 引擎 =================
OCR_ENGINE = None

def init_ppocrv6():
    global OCR_ENGINE
    if OCR_ENGINE is not None:
        return OCR_ENGINE

    print("[Info] 正在加载并初始化 PP-OCRv6 Tiny 1.5m 模型...")
    try:
        # 自动下载/获取模型本地路径
        det_dir = snapshot_download('PaddlePaddle/PP-OCRv6_tiny_det_onnx')
        rec_dir = snapshot_download('PaddlePaddle/PP-OCRv6_tiny_rec_onnx')
        
        det_path = os.path.join(det_dir, 'inference.onnx')
        rec_path = os.path.join(rec_dir, 'inference.onnx')
        
        # 1. 自动从 inference.yml 中提取正确的 PP-OCRv6 字典并保存为 txt 字典文件
        import yaml
        yml_path = os.path.join(rec_dir, 'inference.yml')
        with open(yml_path, 'r', encoding='utf-8') as f:
            yml_content = yaml.safe_load(f)
        char_list = yml_content['PostProcess']['character_dict']
        
        dict_txt_path = os.path.join(UPLOAD_DIR, 'dict_ppocrv6.txt').replace('\\', '/')
        with open(dict_txt_path, 'w', encoding='utf-8') as df:
            for char in char_list:
                df.write(char + '\n')
        
        # 引入 monkey patch 解决 rapidocr_onnxruntime 对 rec_keys_path 参数传递的 Bug
        import rapidocr_onnxruntime.utils
        orig_update_rec_params = rapidocr_onnxruntime.utils.UpdateParameters.update_rec_params

        def patched_update_rec_params(self, config, rec_dict):
            res_config = orig_update_rec_params(self, config, rec_dict)
            if 'rec_keys_path' in rec_dict:
                res_config['keys_path'] = rec_dict['rec_keys_path']
            return res_config

        rapidocr_onnxruntime.utils.UpdateParameters.update_rec_params = patched_update_rec_params

        # 2. 直接使用官方高阶 API 实例化，传入自定义的模型路径和字典文件
        # 避免由于不同机器/容器中 rapidocr_onnxruntime 版本不同导致没有 init_module 私有属性的 Bug
        OCR_ENGINE = RapidOCR(
            det_model_path=det_path,
            rec_model_path=rec_path,
            rec_keys_path=dict_txt_path,
            use_angle_cls=False
        )
        
        print("[Success] PP-OCRv6 Tiny ONNX 引擎加载成功！")
        return OCR_ENGINE
    except Exception as e:
        print(f"[Warning] 无法加载 PP-OCRv6，回退至默认 OCR 模型: {e}")
        OCR_ENGINE = RapidOCR(use_angle_cls=False)
        return OCR_ENGINE


# ================= 2. 图像裁剪核心算法 (移植自 run.py) =================
TARGET_ASPECT_RATIO = 3.37 / 2.13

def _order_points(pts):
    rect = np.zeros((4, 2), dtype="float32")
    s = pts.sum(axis=1)
    rect[0] = pts[np.argmin(s)]
    rect[2] = pts[np.argmax(s)]
    diff = np.diff(pts, axis=1)
    rect[1] = pts[np.argmin(diff)]
    rect[3] = pts[np.argmax(diff)]
    return rect

def _expand_quad(quad, img_shape, scale_x=1.03, scale_y=1.05):
    if quad is None:
        return None
    pts = np.asarray(quad, dtype="float32")
    center = np.mean(pts, axis=0)
    expanded = pts.copy()
    expanded[:, 0] = center[0] + (expanded[:, 0] - center[0]) * scale_x
    expanded[:, 1] = center[1] + (expanded[:, 1] - center[1]) * scale_y
    h, w = img_shape[:2]
    expanded[:, 0] = np.clip(expanded[:, 0], 0, w - 1)
    expanded[:, 1] = np.clip(expanded[:, 1], 0, h - 1)
    return _order_points(expanded)

def _four_point_transform(image, pts):
    rect = _order_points(pts)
    (tl, tr, br, bl) = rect
    width_a = np.linalg.norm(br - bl)
    width_b = np.linalg.norm(tr - tl)
    max_width = int(max(width_a, width_b))
    height_a = np.linalg.norm(tr - br)
    height_b = np.linalg.norm(tl - bl)
    max_height = int(max(height_a, height_b))
    if max_width <= 0 or max_height <= 0:
        return None
    dst = np.array(
        [[0, 0], [max_width - 1, 0], [max_width - 1, max_height - 1], [0, max_height - 1]],
        dtype="float32",
    )
    m = cv2.getPerspectiveTransform(rect, dst)
    return cv2.warpPerspective(image, m, (max_width, max_height), borderValue=(255, 255, 255))



def _scan_content_bound(signal, threshold, max_trim, forward=True, min_run=3):
    length = len(signal)
    if length == 0:
        return 0 if forward else -1
    start = 0 if forward else length - 1
    stop = min(max_trim, length - 1)
    step = 1 if forward else -1
    run = 0
    last_idx = start
    for offset in range(stop + 1):
        idx = start + offset * step
        if float(signal[idx]) >= threshold:
            run += 1
            if run >= min_run:
                return idx - (min_run - 1) * step
        else:
            run = 0
        last_idx = idx
    return last_idx

def trim_document_borders(img, max_trim_ratio=0.005):
    if img is None or img.size == 0:
        return img
    if len(img.shape) == 3:
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    else:
        gray = img.copy()
    h, w = gray.shape[:2]
    if h < 20 or w < 20:
        return img
    gray = cv2.GaussianBlur(gray, (5, 5), 0)
    grad_x = cv2.Sobel(gray, cv2.CV_32F, 1, 0, ksize=3)
    grad_y = cv2.Sobel(gray, cv2.CV_32F, 0, 1, ksize=3)
    grad = cv2.magnitude(grad_x, grad_y)
    band = max(3, int(min(h, w) * 0.02))
    bg_samples = np.concatenate([
        gray[:band, :].reshape(-1),
        gray[-band:, :].reshape(-1),
        gray[:, :band].reshape(-1),
        gray[:, -band:].reshape(-1),
    ])
    bg_level = float(np.median(bg_samples))
    diff = cv2.absdiff(gray, np.full_like(gray, int(round(bg_level))))
    col_signal = np.percentile(diff, 80, axis=0) + np.percentile(grad, 80, axis=0) * 0.6
    row_signal = np.percentile(diff, 80, axis=1) + np.percentile(grad, 80, axis=1) * 0.6
    col_threshold = max(float(np.median(col_signal) + np.std(col_signal) * 1.2), 8.0)
    row_threshold = max(float(np.median(row_signal) + np.std(row_signal) * 1.2), 8.0)
    max_trim_x = max(2, int(w * max_trim_ratio))
    max_trim_y = max(2, int(h * max_trim_ratio))
    left = _scan_content_bound(col_signal, col_threshold, max_trim_x, forward=True)
    right = _scan_content_bound(col_signal, col_threshold, max_trim_x, forward=False)
    top = _scan_content_bound(row_signal, row_threshold, max_trim_y, forward=True)
    bottom = _scan_content_bound(row_signal, row_threshold, max_trim_y, forward=False)
    if right - left < w * 0.5 or bottom - top < h * 0.5:
        return img
    if left <= 0 and top <= 0 and right >= w - 1 and bottom >= h - 1:
        return img
    return img[top:bottom + 1, left:right + 1]

# 方向矫正函数统一从 app.orient 导入，避免重复代码
from app.orient import orient_card_result as _orient_card_result

def pad_image_to_ratio(img):
    """用留白把图片调整到身份证比例，绝不再次裁掉已有内容。"""
    h, w = img.shape[:2]
    if h == 0 or w == 0:
        return img
    current_ratio = w / h
    # 已经接近目标比例，直接返回
    if 1.5 < current_ratio < 1.68:
        return img

    if current_ratio < TARGET_ASPECT_RATIO:
        new_w = int(round(h * TARGET_ASPECT_RATIO))
        total_pad = max(0, new_w - w)
        left = total_pad // 2
        right = total_pad - left
        return cv2.copyMakeBorder(img, 0, 0, left, right, cv2.BORDER_CONSTANT, value=[255, 255, 255])
    else:
        new_h = int(round(w / TARGET_ASPECT_RATIO))
        total_pad = max(0, new_h - h)
        top = total_pad // 2
        bottom = total_pad - top
        return cv2.copyMakeBorder(img, top, bottom, 0, 0, cv2.BORDER_CONSTANT, value=[255, 255, 255])

def _find_best_card_quad(contours, img_area, min_area_ratio, max_area_ratio, min_ratio, max_ratio, target_ratio, strict_ratio=True):
    best = None
    best_score = 0.0
    for c in contours:
        area = cv2.contourArea(c)
        if area < img_area * min_area_ratio or area > img_area * max_area_ratio:
            continue
        hull = cv2.convexHull(c)
        peri = cv2.arcLength(hull, True)
        if peri == 0:
            continue
        approx = cv2.approxPolyDP(hull, 0.015 * peri, True)
        if len(approx) != 4:
            approx = cv2.approxPolyDP(hull, 0.03 * peri, True)
        use_min_rect = True
        if len(approx) == 4:
            candidate_approx = approx.reshape(4, 2)
            is_rect = True
            for i in range(4):
                pt1, pt2, pt3 = candidate_approx[i], candidate_approx[(i+1)%4], candidate_approx[(i+2)%4]
                v1, v2 = pt1 - pt2, pt3 - pt2
                cos_a = np.dot(v1, v2) / (np.linalg.norm(v1) * np.linalg.norm(v2) + 1e-5)
                angle = np.degrees(np.arccos(np.clip(cos_a, -1.0, 1.0)))
                if abs(angle - 90) > 12:
                    is_rect = False
                    break
            if is_rect:
                candidate = candidate_approx
                use_min_rect = False
        if use_min_rect:
            min_rect = cv2.minAreaRect(hull)
            candidate = cv2.boxPoints(min_rect)
        rect = _order_points(candidate.astype("float32"))
        rw = np.linalg.norm(rect[1] - rect[0])
        rh = np.linalg.norm(rect[3] - rect[0])
        if rh <= 0 or rw <= 0:
            continue
        ratio = rw / rh if rw > rh else rh / rw
        if strict_ratio and (ratio < min_ratio or ratio > max_ratio):
            continue
        ratio_score = max(0.0, 1.0 - abs(ratio - target_ratio) / target_ratio)
        rect_area = rw * rh
        rectangularity = area / rect_area if rect_area > 0 else 0
        area_ratio = area / img_area
        score = ratio_score * 1000.0 + area_ratio * 600.0 + rectangularity * 300.0
        if score > best_score:
            best_score = score
            best = rect
    return best

def crop_by_card_contour(img, is_special_cert=False):
    if img is None or img.size == 0:
        return None
    if len(img.shape) == 3:
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    else:
        gray = img.copy()
    h, w = gray.shape[:2]
    img_area = h * w
    if img_area == 0:
        return None
    min_area_ratio = 0.05
    max_area_ratio = 0.99
    min_ratio, max_ratio = 1.35, 1.85

    best = None
    gray_b = cv2.bilateralFilter(gray, 11, 17, 17)
    for low, high in [(30, 200), (20, 100), (10, 50)]:
        edged = cv2.Canny(gray_b, low, high)
        kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (5, 5))
        edged = cv2.morphologyEx(edged, cv2.MORPH_CLOSE, kernel, iterations=2)
        contours, _ = cv2.findContours(edged, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
        best = _find_best_card_quad(
            contours, img_area, min_area_ratio, max_area_ratio, min_ratio, max_ratio, TARGET_ASPECT_RATIO, strict_ratio=True
        )
        if best is not None:
            break

    if best is None:
        for block_size in [51, 71, 91]:
            thresh = cv2.adaptiveThreshold(gray_b, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                                          cv2.THRESH_BINARY_INV, block_size, 10)
            kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (5, 5))
            thresh = cv2.morphologyEx(thresh, cv2.MORPH_CLOSE, kernel, iterations=3)
            contours, _ = cv2.findContours(thresh, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
            best = _find_best_card_quad(
                contours, img_area, min_area_ratio, max_area_ratio, min_ratio, max_ratio, TARGET_ASPECT_RATIO, strict_ratio=False
            )
            if best is not None:
                break

    if best is None:
        return None

    # 轮廓通常落在卡片印刷区内侧，向外扩少量可避免切掉实体卡边和底部号码。
    warped = _four_point_transform(img, _expand_quad(best, img.shape, scale_x=1.05, scale_y=1.08))
    if warped is None:
        return None
    return warped  # 方向已由 _orient_by_ocr 校正，不在此处旋转

def _detect_id_card_body_quad(img):
    if img is None or img.size == 0 or len(img.shape) != 3:
        return None
    h, w = img.shape[:2]
    img_area = float(h * w)
    if img_area <= 0:
        return None
    hsv = cv2.cvtColor(img, cv2.COLOR_BGR2HSV)
    lab = cv2.cvtColor(img, cv2.COLOR_BGR2LAB)
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    s = hsv[:, :, 1]
    v = hsv[:, :, 2]
    l = lab[:, :, 0]
    a = lab[:, :, 1]
    b = lab[:, :, 2]
    low_sat = cv2.inRange(s, 0, 120)
    bright = cv2.inRange(v, 90, 255)
    warm_a = cv2.inRange(a, 110, 165)
    warm_b = cv2.inRange(b, 110, 180)
    soft_l = cv2.inRange(l, 90, 250)
    mask = cv2.bitwise_and(low_sat, bright)
    mask = cv2.bitwise_and(mask, soft_l)
    warm_mask = cv2.bitwise_and(warm_a, warm_b)
    mask = cv2.bitwise_or(mask, warm_mask)
    kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (7, 7))
    mask = cv2.morphologyEx(mask, cv2.MORPH_CLOSE, kernel, iterations=2)
    mask = cv2.morphologyEx(mask, cv2.MORPH_OPEN, kernel, iterations=1)
    edges = cv2.Canny(gray, 40, 140)
    edges = cv2.dilate(edges, cv2.getStructuringElement(cv2.MORPH_RECT, (3, 3)), iterations=1)
    mask = cv2.bitwise_or(mask, edges)
    mask = cv2.morphologyEx(mask, cv2.MORPH_CLOSE, kernel, iterations=2)
    contours, _ = cv2.findContours(mask, cv2.RETR_LIST, cv2.CHAIN_APPROX_SIMPLE)
    quad = _find_best_card_quad(
        contours,
        img_area,
        min_area_ratio=0.08,
        max_area_ratio=0.95,
        min_ratio=1.35,
        max_ratio=1.90,
        target_ratio=TARGET_ASPECT_RATIO,
        strict_ratio=False,
    )
    if quad is None:
        return None
    rect = _order_points(np.asarray(quad, dtype="float32"))
    rw = np.linalg.norm(rect[1] - rect[0])
    rh = np.linalg.norm(rect[3] - rect[0])
    if rw <= 0 or rh <= 0:
        return None
    ratio = rw / rh if rw >= rh else rh / rw
    if abs(ratio - TARGET_ASPECT_RATIO) > 0.45:
        return None
    fill_ratio = _quad_polygon_area(rect) / img_area
    if fill_ratio < 0.08 or fill_ratio > 0.95:
        return None
    return rect

def _quad_polygon_area(quad):
    try:
        return float(abs(cv2.contourArea(np.asarray(quad, dtype="float32"))))
    except Exception:
        return 0.0

def _crop_id_card_image(img):
    if img is None:
        return img
    try:
        quad = _detect_id_card_body_quad(img)
        if quad is not None:
            body = _four_point_transform(img, _expand_quad(quad, img.shape, scale_x=1.03, scale_y=1.05))
            if body is not None and body.size != 0:
                # 方向已由 _orient_by_ocr 校正，不在此处旋转
                body = trim_document_borders(body, max_trim_ratio=0.005)
                return pad_image_to_ratio(body)
        contour_result = crop_by_card_contour(img, is_special_cert=False)
        if contour_result is not None:
            ch, cw = contour_result.shape[:2]
            ih, iw = img.shape[:2]
            if ch * cw < ih * iw * 0.95:
                # 方向已由 _orient_by_ocr 校正，不在此处旋转
                contour_result = trim_document_borders(contour_result, max_trim_ratio=0.005)
                return pad_image_to_ratio(contour_result)
        work = img.copy()
        # 方向已由 _orient_by_ocr 校正，不在此处旋转
        work = trim_document_borders(work, max_trim_ratio=0.01)
        return pad_image_to_ratio(work)
    except Exception:
        return img

# ================= 3. 信息智能解析 =================
def clean_name_string(name_str):
    name_str = str(name_str).replace(" ", "").replace("　", "")
    if "葡" in name_str and "葡萄" not in name_str:
        name_str = name_str.replace("葡", "蔺")
    return name_str



def extract_nation(full_text):
    # D5: 统一引用 app.postprocess.NATIONS，避免56个民族列表重复维护
    for n in NATIONS:
        if n in full_text:
            return n + "族"
    return "汉族"

def smart_extract_info(ocr_results):
    if not ocr_results:
        return "", "", ""
    lines_clean = [line[1].replace(" ", "") for line in ocr_results]
    full_text_clean = "".join(lines_clean).upper()
    id_number = ""
    id_match = re.search(r'([1-9]\d{5}[12]\d{3}[01]\d[0123]\d{4}[\dXx])', full_text_clean)
    if not id_match:
        corrected = full_text_clean.replace('O', '0').replace('Z', '2').replace('I', '1').replace('S', '5')
        id_match = re.search(r'([1-9]\d{5}[12]\d{3}[01]\d[0123]\d{4}[\dXx])', corrected)
    if id_match:
        id_number = id_match.group(1).upper()

    name = ""
    
    # 策略 1: 直接匹配包含“姓名”的文本行
    for i, line in enumerate(ocr_results):
        text = str(line[1]).replace(" ", "").replace(":", "").replace("：", "")
        if "姓名" in text:
            # 如果这行字不止“姓名”两个字，例如“姓名张三”
            if len(text) > 2:
                candidate = text.replace("姓名", "")
                candidate = re.sub(r'[^\u4e00-\u9fa5]', '', candidate)
                if 2 <= len(candidate) <= 4 and not any(k in candidate for k in ["性别", "民族", "出生", "住址", "号码"]):
                    name = candidate
                    break
            # 如果这行只有“姓名”，找它同行的右侧框，或者下一行
            else:
                box_name = np.array(line[0])
                name_y = np.mean(box_name[:, 1])
                name_x_max = np.max(box_name[:, 0])
                
                # 寻找右侧同行且最近的框
                best_right_cand = ""
                best_right_dist = 9999
                for r_line in ocr_results:
                    r_text = str(r_line[1]).replace(" ", "")
                    if r_text == text:
                        continue
                    r_box = np.array(r_line[0])
                    r_y = np.mean(r_box[:, 1])
                    r_x_min = np.min(r_box[:, 0])
                    
                    # 同行（y轴差值小于15）且在右边
                    if abs(r_y - name_y) < 15 and r_x_min > name_x_max - 5:
                        dist = r_x_min - name_x_max
                        if dist < best_right_dist:
                            best_right_dist = dist
                            best_right_cand = re.sub(r'[^\u4e00-\u9fa5]', '', r_text)
                            
                if best_right_cand and 2 <= len(best_right_cand) <= 4 and not any(k in best_right_cand for k in ["性别", "民族", "出生", "住址"]):
                    name = best_right_cand
                    break
                    
                # 或者是下一行
                if i + 1 < len(ocr_results):
                    next_text = str(ocr_results[i + 1][1]).replace(" ", "")
                    cleaned = re.sub(r'[^\u4e00-\u9fa5]', '', next_text)
                    if 2 <= len(cleaned) <= 4 and not any(k in cleaned for k in ["性别", "民族", "出生", "住址"]):
                        name = cleaned
                        break

    # 策略 2: 基于“性别”或“民族”的 y 坐标向上寻找
    if not name or len(name) < 2:
        gender_box_y = -1
        for line in ocr_results:
            text = str(line[1]).replace(" ", "")
            if "性别" in text or "民族" in text:
                box = np.array(line[0])
                gender_box_y = np.mean(box[:, 1])
                break
                
        if gender_box_y != -1:
            best_name_candidate = ""
            best_name_y_diff = 9999
            for line in ocr_results:
                text = str(line[1]).replace(" ", "")
                box = np.array(line[0])
                cy = np.mean(box[:, 1])
                
                # 高度差必须大于 15 像素，避免把同行的“民族汉”误认作姓名
                if 15 < (gender_box_y - cy) < 200: 
                    clean_text = text.replace("姓名", "").replace("名", "")
                    clean_text = re.sub(r'[^\u4e00-\u9fa5]', '', clean_text)
                    if 2 <= len(clean_text) <= 4 and not any(k in clean_text for k in ["性别", "民族", "出生", "住址"]):
                        if (gender_box_y - cy) < best_name_y_diff:
                            best_name_y_diff = gender_box_y - cy
                            best_name_candidate = clean_text
            if best_name_candidate:
                name = best_name_candidate

    # 策略 3: 兜底方案，取前几行
    if not name or len(name) < 2:
        for line_text in lines_clean[:3]:
            cleaned = re.sub(r'[^\u4e00-\u9fa5]', '', line_text).replace("姓名", "", 1)
            if 2 <= len(cleaned) <= 4 and not any(k in cleaned for k in ["性别", "民族", "出生", "住址", "号码", "身份"]):
                name = cleaned
                break

    # 终极黑名单过滤与清洗
    name = clean_name_string(name)
    invalid_keywords = ["性别", "民族", "出生", "住址", "号码", "公民", "身份", "汉族", "别男", "别女"]
    if any(k in name for k in invalid_keywords):
        name = ""
    if name in ["汉", "男", "女", "汉族", "族", "男新", "女新", "姓名", "姓名名", "名"]:
        name = ""

    nation = extract_nation(full_text_clean)
    return name[:5] if len(name) >= 2 else "", id_number, nation

def _read_and_auto_orient(image_path):
    """用 PIL 读取图片并自动应用 EXIF 方向信息（修正手机竖拍/横拍旋转），返回 OpenCV BGR 图像"""
    try:
        pil_img = Image.open(image_path)
        pil_img = ImageOps.exif_transpose(pil_img)
        if pil_img.mode != 'RGB':
            pil_img = pil_img.convert('RGB')
        return cv2.cvtColor(np.array(pil_img), cv2.COLOR_RGB2BGR)
    except Exception:
        # fallback: 直接用 OpenCV 读取（不处理 EXIF）
        return cv2.imdecode(np.fromfile(image_path, dtype=np.uint8), cv2.IMREAD_COLOR)

def _crop_by_ocr_boxes(img, ocr_results):
    """用身份证上的锚点文字精确定位卡面边界，然后裁切。
    
    标准身份证布局 (85.6mm × 54mm, 比例 ≈ 1.585):
    - "姓名" / "性别" / "住址" 等标签左起: 约 6.0% 卡宽
    - "姓名"标签顶部 y: 约 11.1% 卡高
    - 身份证号区域:
      - 包含"公民身份号码"标签时: 占卡宽约 [6.0%, 90.0%]
      - 仅数字时: 占卡宽约 [38.0%, 90.0%]
      - 底部 y: 约 93.5% 卡高
    """
    if not ocr_results or len(ocr_results) < 2:
        return img

    h, w = img.shape[:2]

    # ===== 寻找锚点 =====
    xingming_box = None
    id_number_box = None
    id_number_text = ""

    for line in ocr_results:
        text = str(line[1]).replace(" ", "").replace("　", "")
        box = np.array(line[0])

        # 锚点A: "姓名" 标签
        if "姓名" in text and xingming_box is None:
            xingming_box = box

        # 锚点B: 18位身份证号
        cleaned = text.upper().replace('O', '0').replace('Z', '2').replace('I', '1').replace('S', '5')
        if re.search(r'[1-9]\d{5}[12]\d{3}[01]\d[0123]\d{4}[\dXx]', cleaned):
            id_number_box = box
            id_number_text = text

    # ===== 基于锚点进行精确定位 =====
    if id_number_box is not None:
        id_left_x = float(np.min(id_number_box[:, 0]))
        id_right_x = float(np.max(id_number_box[:, 0]))
        id_bottom_y = float(np.max(id_number_box[:, 1]))

        # 如果同时找到了姓名和身份证号，则使用两者作为卡宽的首尾定位锚点
        # "姓名"左侧是 6.0% 卡宽，身份证号识别框右端是 90.0% 卡宽。双锚点跨度为 84.0%
        if xingming_box is not None:
            xm_left_x = float(np.min(xingming_box[:, 0]))
            xm_top_y = float(np.min(xingming_box[:, 1]))
            
            card_w_est = (id_right_x - xm_left_x) / 0.84
            card_left = xm_left_x - card_w_est * 0.06
            card_right = card_left + card_w_est
            
            anchor_span_y = id_bottom_y - xm_top_y
            if anchor_span_y > h * 0.1:
                card_h_est = anchor_span_y / 0.824
                card_top = xm_top_y - card_h_est * 0.111
                card_bottom = card_top + card_h_est
            else:
                card_h_est = card_w_est / 1.585
                card_bottom = id_bottom_y + card_h_est * 0.065
                card_top = card_bottom - card_h_est
        else:
            # 仅有身份证号锚点时，判断是否包含中文标签
            has_label = any(c in id_number_text for c in ["公", "民", "身", "份", "号", "码"]) or bool(re.search(r'[\u4e00-\u9fa5]', id_number_text))
            if has_label:
                box_left_ratio = 0.06
            else:
                box_left_ratio = 0.38
            box_right_ratio = 0.90
            box_w_ratio = box_right_ratio - box_left_ratio
            
            card_w_est = (id_right_x - id_left_x) / box_w_ratio
            card_left = id_left_x - card_w_est * box_left_ratio
            card_right = card_left + card_w_est

            card_h_est = card_w_est / 1.585
            card_bottom = id_bottom_y + card_h_est * 0.065
            card_top = card_bottom - card_h_est

        # OCR框落在文字内侧，按估算卡宽/卡高再向外保留安全边距。
        # 手机拍屏时卡片边框常被屏幕横纹削弱，宁可多留少量背景，也不能切掉证件内容。
        safety_x = card_w_est * 0.03
        safety_y = card_h_est * 0.04
        y1 = int(max(0, card_top - safety_y))
        y2 = int(min(h, card_bottom + safety_y))
        x1 = int(max(0, card_left - safety_x))
        x2 = int(min(w, card_right + safety_x))

        crop_area = (x2 - x1) * (y2 - y1)
        if 0 < crop_area < h * w * 0.95 and (x2 - x1) > w * 0.18 and (y2 - y1) > h * 0.08:
            cropped = img[y1:y2, x1:x2]
            if cropped.size > 0:
                print(f"[OCR-Debug] 锚点精确裁切: left={x1} right={x2} top={y1} bottom={y2} 宽={x2-x1} 高={y2-y1}")
                return cropped

    # ===== 回退：全部OCR框 + IQR过滤 =====
    box_centers = []
    for line in ocr_results:
        box = np.array(line[0])
        cy = float(np.mean(box[:, 1]))
        cx = float(np.mean(box[:, 0]))
        box_centers.append((cx, cy, box))

    ys = np.array([c[1] for c in box_centers])
    q1_y, q3_y = np.percentile(ys, 25), np.percentile(ys, 75)
    iqr_y = q3_y - q1_y
    fence = max(iqr_y * 1.5, 20)
    y_lo, y_hi = q1_y - fence, q3_y + fence

    xs = np.array([c[0] for c in box_centers])
    q1_x, q3_x = np.percentile(xs, 25), np.percentile(xs, 75)
    iqr_x = q3_x - q1_x
    fence_x = max(iqr_x * 1.5, 20)
    x_lo, x_hi = q1_x - fence_x, q3_x + fence_x

    filtered_points = []
    for cx, cy, box in box_centers:
        if y_lo <= cy <= y_hi and x_lo <= cx <= x_hi:
            filtered_points.extend(box.tolist())

    if len(filtered_points) < 8:
        filtered_points = []
        for _, _, box in box_centers:
            filtered_points.extend(box.tolist())

    all_points = np.array(filtered_points)
    min_x = int(np.min(all_points[:, 0]))
    max_x = int(np.max(all_points[:, 0]))
    min_y = int(np.min(all_points[:, 1]))
    max_y = int(np.max(all_points[:, 1]))

    text_w = max_x - min_x
    text_h = max_y - min_y

    if text_w < w * 0.15 or text_h < h * 0.15:
        return img

    top_margin = text_h * 0.135
    bottom_margin = text_h * 0.091
    left_margin = text_w * 0.067
    right_margin = text_w * 0.08

    y1 = max(0, int(min_y - top_margin))
    y2 = min(h, int(max_y + bottom_margin))
    x1 = max(0, int(min_x - left_margin))
    x2 = min(w, int(max_x + right_margin))

    crop_area = (x2 - x1) * (y2 - y1)
    if crop_area > h * w * 0.92:
        return img

    cropped = img[y1:y2, x1:x2]
    if cropped.size == 0:
        return img
    print(f"[OCR-Debug] IQR回退裁切: 文字区{text_w}x{text_h} 边缘({min_x},{min_y})-({max_x},{max_y}) → 裁切{x2-x1}x{y2-y1}")
    return cropped

def _ocr_quality(ocr_results):
    """按身份证字段完整度和文字方向评分，避免把旋转后的屏幕文字当成有效卡面。"""
    if not ocr_results:
        return 0.0, "", "", ""
    name, id_card, nation = smart_extract_info(ocr_results)
    valid_id = bool(re.fullmatch(r'[1-9]\d{16}[\dX]', id_card or ""))
    score = 180.0 if valid_id else 0.0
    if name:
        score += 55.0

    keywords = ("姓名", "性别", "民族", "出生", "住址", "公民身份号码")
    texts = [str(line[1]).replace(" ", "") for line in ocr_results if len(line) >= 2]
    score += min(6, sum(1 for key in keywords if any(key in text for text in texts))) * 8.0

    confidences = []
    horizontal_lines = 0
    measured_lines = 0
    for line in ocr_results:
        if len(line) >= 3:
            try:
                confidences.append(float(line[2]))
            except (TypeError, ValueError):
                pass
        if len(line) < 2:
            continue
        box = np.asarray(line[0], dtype="float32")
        if box.shape != (4, 2):
            continue
        box_w = float(np.max(box[:, 0]) - np.min(box[:, 0]))
        box_h = float(np.max(box[:, 1]) - np.min(box[:, 1]))
        if len(str(line[1]).strip()) >= 2 and box_w > 0 and box_h > 0:
            measured_lines += 1
            if box_w >= box_h * 1.35:
                horizontal_lines += 1

        cleaned = str(line[1]).upper().replace(" ", "").replace('O', '0').replace('Z', '2').replace('I', '1').replace('S', '5')
        if re.search(r'[1-9]\d{5}[12]\d{3}[01]\d[0123]\d{4}[\dX]', cleaned):
            if box_w >= box_h * 2.0:
                score += 70.0
            elif box_h >= box_w * 2.0:
                score -= 90.0

    if measured_lines:
        score += (horizontal_lines / measured_lines) * 25.0
    if confidences:
        score += max(0.0, min(1.0, float(np.mean(confidences)))) * 20.0
    return score, name, id_card, nation


def _run_ocr(engine, image):
    if image is None or image.size == 0:
        return None, (0.0, "", "", "")
    try:
        results, _ = engine(image)
        return results, _ocr_quality(results)
    except Exception as exc:
        print(f"[OCR-Debug] OCR候选识别失败: {type(exc).__name__}")
        return None, (0.0, "", "", "")


def _reduce_screen_pattern(image):
    """轻度抑制拍摄屏幕产生的横纹/摩尔纹，仅用于OCR，不改动最终保存图。"""
    if image is None or image.size == 0:
        return image
    softened = cv2.GaussianBlur(image, (3, 3), 0.75)
    lab = cv2.cvtColor(softened, cv2.COLOR_BGR2LAB)
    l_chan, a_chan, b_chan = cv2.split(lab)
    l_chan = cv2.createCLAHE(clipLimit=1.6, tileGridSize=(8, 8)).apply(l_chan)
    return cv2.cvtColor(cv2.merge((l_chan, a_chan, b_chan)), cv2.COLOR_LAB2BGR)


def _select_input_orientation(engine, image):
    """OCR选优方向；不再因为照片是竖图就无条件旋转。"""
    methods = ("original", "rotate_cw", "rotate_ccw", "rotate_180")
    best = None
    for index, method in enumerate(methods):
        if method == "original":
            candidate = image
        elif method == "rotate_cw":
            candidate = cv2.rotate(image, cv2.ROTATE_90_CLOCKWISE)
        elif method == "rotate_ccw":
            candidate = cv2.rotate(image, cv2.ROTATE_90_COUNTERCLOCKWISE)
        else:
            candidate = cv2.rotate(image, cv2.ROTATE_180)
        results, metrics = _run_ocr(engine, candidate)
        item = (metrics[0], method, candidate, results, metrics)
        if best is None or item[0] > best[0]:
            best = item
        # 正向原图已完整识别时，不再做额外三次OCR。
        if index == 0 and metrics[2] and metrics[1] and metrics[0] >= 260:
            break

    if best is None:
        return image, None, (0.0, "", "", ""), "original"

    # 信息仍不完整时再尝试一次屏幕纹理抑制，框坐标与原图保持一致。
    if not best[4][2] or not best[4][1]:
        filtered = _reduce_screen_pattern(best[2])
        filtered_results, filtered_metrics = _run_ocr(engine, filtered)
        if filtered_metrics[0] > best[0]:
            best = (filtered_metrics[0], best[1] + "_screen_filter", best[2], filtered_results, filtered_metrics)
    return best[2], best[3], best[4], best[1]


def _prepare_card_candidate(candidate):
    if candidate is None or candidate.size == 0:
        return None
    h, w = candidate.shape[:2]
    if h > w and 1.35 <= h / max(w, 1) <= 1.90:
        candidate = cv2.rotate(candidate, cv2.ROTATE_90_CLOCKWISE)
    if candidate.shape[1] >= candidate.shape[0]:
        candidate = _orient_card_result(candidate)
    return pad_image_to_ratio(candidate)


def _is_real_crop(candidate, source):
    if candidate is None or candidate.size == 0:
        return False
    ch, cw = candidate.shape[:2]
    sh, sw = source.shape[:2]
    if ch < 80 or cw < 120:
        return False
    return ch * cw < sh * sw * 0.97


def _enhance_document_edges(image, screen_mode=False):
    """提高文档预处理器的边缘输入质量，不改变原始 OCR 主输入。"""
    if image is None or image.size == 0:
        return image
    source = _reduce_screen_pattern(image) if screen_mode else image
    # 在 LAB 的亮度通道上做温和局部对比度增强，保留证件颜色和人像，
    # 再做轻度锐化，使被压缩、拍屏或轻微虚焦的卡边更容易被 Canny 找到。
    lab = cv2.cvtColor(source, cv2.COLOR_BGR2LAB)
    l_chan, a_chan, b_chan = cv2.split(lab)
    l_chan = cv2.createCLAHE(
        clipLimit=2.0 if screen_mode else 1.7,
        tileGridSize=(8, 8)
    ).apply(l_chan)
    enhanced = cv2.cvtColor(cv2.merge((l_chan, a_chan, b_chan)), cv2.COLOR_LAB2BGR)
    softened = cv2.GaussianBlur(enhanced, (0, 0), 1.1)
    return cv2.addWeighted(enhanced, 1.20, softened, -0.20, 0)


def _document_preprocessor_profile(profile):
    """针对满屏证件、普通拍摄和屏幕拍摄提供不同的轮廓阈值。"""
    profiles = {
        # 默认库在 1000px 下处理，满屏或高分辨率证件的边缘会被明显缩弱。
        "standard": {
            "max_proc_dim": 1600,
            "canny_threshold1": 35,
            "canny_threshold2": 120,
            "contour_epsilon_coef": 0.024,
            "min_page_area_ratio": 0.16,
            "min_rectangularity": 0.62,
            "screen_mode": False,
        },
        # 对轻微虚焦、压缩和卡片占画面较小的拍照，放宽边缘与矩形门槛。
        "edge_enhanced": {
            "max_proc_dim": 1600,
            "canny_threshold1": 22,
            "canny_threshold2": 96,
            "contour_epsilon_coef": 0.028,
            "min_page_area_ratio": 0.10,
            "min_rectangularity": 0.56,
            "screen_mode": False,
        },
        # 拍摄显示器时先抑制横纹/摩尔纹，再使用较低的边缘阈值。
        "screen_enhanced": {
            "max_proc_dim": 1600,
            "canny_threshold1": 16,
            "canny_threshold2": 78,
            "contour_epsilon_coef": 0.030,
            "min_page_area_ratio": 0.09,
            "min_rectangularity": 0.52,
            "screen_mode": True,
        },
    }
    return profiles[profile]


def _document_preprocessor_candidate(image, profile="standard"):
    """使用开源文档裁切器生成候选卡面，并按输入场景调节检测参数。"""
    try:
        from document_preprocessor import DocumentPreprocessor
        from document_preprocessor.core import PreprocessorConfig
        import PIL.Image

        settings = _document_preprocessor_profile(profile)
        source = (
            _enhance_document_edges(image, screen_mode=settings["screen_mode"])
            if profile != "standard" else image
        )
        config = PreprocessorConfig(
            max_proc_dim=settings["max_proc_dim"],
            canny_threshold1=settings["canny_threshold1"],
            canny_threshold2=settings["canny_threshold2"],
            contour_epsilon_coef=settings["contour_epsilon_coef"],
            min_page_area_ratio=settings["min_page_area_ratio"],
            min_rectangularity=settings["min_rectangularity"],
        )
        preprocessor = DocumentPreprocessor(config=config)
        pil_img = PIL.Image.fromarray(cv2.cvtColor(source, cv2.COLOR_BGR2RGB))
        pil_warped = preprocessor.detect_and_warp_document(pil_img)
        candidate = cv2.cvtColor(np.asarray(pil_warped), cv2.COLOR_RGB2BGR)
        if _is_real_crop(candidate, image):
            return candidate
        # 库在找不到四角时会原样返回输入图；这不是有效裁切，继续尝试
        # 更适配的第一道 profile，而不是提前把“原图”标记为成功。
        print(f"[OCR-Debug] 文档预处理未找到有效卡边: {profile}")
        return None
    except Exception as exc:
        print(f"[OCR-Debug] document-preprocessor不可用: {profile}, {type(exc).__name__}")
        return None


def _candidate_score(ocr_score, image):
    h, w = image.shape[:2]
    ratio = w / max(h, 1)
    aspect_bonus = max(0.0, 55.0 - abs(ratio - TARGET_ASPECT_RATIO) * 120.0)
    return ocr_score + aspect_bonus


def _is_valid_id_card_number(id_card):
    """验证 18 位居民身份证格式和校验位，避免误识别文本提前结束回退。"""
    value = str(id_card or "").strip().upper()
    if not re.fullmatch(r"[1-9]\d{16}[\dX]", value):
        return False
    weights = (7, 9, 10, 5, 8, 4, 2, 1, 6, 3, 7, 9, 10, 5, 8, 4, 2)
    check_codes = "10X98765432"
    checksum = sum(int(value[index]) * weights[index] for index in range(17)) % 11
    return value[-1] == check_codes[checksum]


def _has_complete_identity(metrics):
    """当前阶段是否已完整识别姓名和可校验的身份证号。"""
    _, name, id_card, _ = metrics
    clean_name = re.sub(r"[\s·・]", "", str(name or ""))
    return len(clean_name) >= 2 and _is_valid_id_card_number(id_card)


def _is_structurally_safe_crop(candidate, source):
    """防止信息虽然被 OCR 读到，但候选图把卡片边缘或主体切掉。"""
    if not _is_real_crop(candidate, source):
        return False
    height, width = candidate.shape[:2]
    long_edge, short_edge = max(width, height), max(1, min(width, height))
    ratio = long_edge / short_edge
    source_area = max(source.shape[0] * source.shape[1], 1)
    crop_area = height * width
    return 1.28 <= ratio <= 2.00 and crop_area >= source_area * 0.035


def _evaluate_crop_candidate(engine, method, candidate, source):
    """统一做结构检查、方向整理和 OCR 完整度评分。"""
    if not _is_structurally_safe_crop(candidate, source):
        print(f"[OCR-Debug] 裁切候选丢弃: {method}, 原因: 结构不可靠")
        return None
    prepared = _prepare_card_candidate(candidate)
    _, metrics = _run_ocr(engine, prepared)
    score = _candidate_score(metrics[0], prepared)
    complete = _has_complete_identity(metrics)
    print(
        f"[OCR-Debug] 裁切候选: {method}, 尺寸: {prepared.shape[1]}x{prepared.shape[0]}, "
        f"质量分: {score:.1f}, 姓名身份证完整: {'是' if complete else '否'}"
    )
    return {
        "method": method,
        "image": prepared,
        "metrics": metrics,
        "score": score,
        "complete": complete,
    }


def ocr_idcard_process(image_path):
    engine = init_ppocrv6()
    source = _read_and_auto_orient(image_path)
    if source is None:
        raise ValueError("无法解析身份证图片。")
    print(f"[OCR-Debug] 原图尺寸: {source.shape[1]}x{source.shape[0]}")

    # 先用OCR判断卡面真实方向。手机竖拍屏幕里的横向身份证会保持竖图，不会再被误转90度。
    oriented, initial_ocr, initial_metrics, orientation_method = _select_input_orientation(engine, source)
    best_info_score = initial_metrics[0]
    name, id_card, nation = initial_metrics[1], initial_metrics[2], initial_metrics[3]
    print(f"[OCR-Debug] 输入方向: {orientation_method}, OCR框: {len(initial_ocr or [])}, 质量分: {initial_metrics[0]:.1f}")

    # 按人工确认的顺序串行回退：开源文档裁切器 → OpenCV → OCR 锚点。
    # 每一道都必须先重新识别出完整姓名和校验通过的身份证号，才能停止。
    print("[OCR-Debug] 裁切顺序: document_preprocessor -> opencv -> ocr_anchors")
    evaluations = []
    selected = None

    # 第一道：优先开源 document-preprocessor。依次尝试原图、高细节边缘、
    # 拍屏抗横纹三种输入；每种仍必须通过完整身份信息校验。
    for method, profile in (
        ("document_preprocessor", "standard"),
        ("document_preprocessor_edge_enhanced", "edge_enhanced"),
        ("document_preprocessor_screen_enhanced", "screen_enhanced"),
    ):
        candidate = _document_preprocessor_candidate(oriented, profile=profile)
        if candidate is None:
            print(f"[OCR-Debug] 第一道无有效候选: {method}")
            continue
        evaluation = _evaluate_crop_candidate(engine, method, candidate, oriented)
        if evaluation is None:
            continue
        evaluations.append(evaluation)
        if evaluation["complete"]:
            selected = evaluation
            break

    # 第二道：OpenCV 卡面检测，再尝试轮廓检测。
    if selected is None:
        print("[OCR-Debug] 第一道未完整识别，进入第二道 OpenCV")
        opencv_candidates = []
        quad = _detect_id_card_body_quad(oriented)
        if quad is not None:
            body_crop = _four_point_transform(oriented, _expand_quad(quad, oriented.shape, scale_x=1.07, scale_y=1.10))
            opencv_candidates.append(("opencv_card_body", body_crop))
        contour_crop = crop_by_card_contour(oriented, is_special_cert=False)
        opencv_candidates.append(("opencv_card_contour", contour_crop))
        for method, candidate in opencv_candidates:
            if candidate is None:
                print(f"[OCR-Debug] 第二道无有效候选: {method}")
                continue
            evaluation = _evaluate_crop_candidate(engine, method, candidate, oriented)
            if evaluation is None:
                continue
            evaluations.append(evaluation)
            if evaluation["complete"]:
                selected = evaluation
                break

    # 第三道：用 OCR 文字框定位；该方法内部会在锚点不足时使用 IQR 回退。
    if selected is None:
        print("[OCR-Debug] 第二道未完整识别，进入第三道 OCR 文字框")
        if initial_ocr and len(initial_ocr) >= 2:
            anchor_crop = _crop_by_ocr_boxes(oriented, initial_ocr)
            evaluation = _evaluate_crop_candidate(engine, "ocr_anchors", anchor_crop, oriented)
            if evaluation is not None:
                evaluations.append(evaluation)
                if evaluation["complete"]:
                    selected = evaluation
        else:
            print("[OCR-Debug] 第三道无有效 OCR 文字框")

    # 汇总所有候选中的最佳识别字段；仅当裁切图明显提升信息质量时才作为不完整结果保存。
    for evaluation in evaluations:
        metrics = evaluation["metrics"]
        if metrics[0] > best_info_score:
            best_info_score = metrics[0]
            name, id_card, nation = metrics[1], metrics[2], metrics[3]

    if selected is not None:
        img_cropped = selected["image"]
        crop_method = selected["method"]
        _, name, id_card, nation = selected["metrics"]
        best_info_score = selected["metrics"][0]
    else:
        best_partial = max(evaluations, key=lambda item: item["score"], default=None)
        if best_partial is not None and best_partial["metrics"][0] >= initial_metrics[0] + 35:
            img_cropped = best_partial["image"]
            crop_method = best_partial["method"] + "_partial"
            print("[OCR-Debug] 三道均未完整识别，保留信息质量明显更高的安全裁切候选")
        else:
            img_cropped = oriented.copy()
            crop_method = "original_safe"
            print("[OCR-Debug] 三道均未完整识别，保护性保留原图")

    if crop_method != "original_safe" and img_cropped.shape[1] < 856:
        scale = 856.0 / img_cropped.shape[1]
        img_cropped = cv2.resize(img_cropped, None, fx=scale, fy=scale, interpolation=cv2.INTER_CUBIC)
    print(f"[OCR-Debug] 最终裁切: {crop_method}, 尺寸: {img_cropped.shape[1]}x{img_cropped.shape[0]}, 信息分: {best_info_score:.1f}")

    cropped_filename = f"crop_{uuid.uuid4().hex}.jpg"
    cropped_path = os.path.join(TEMP_IDS_DIR, cropped_filename).replace('\\', '/')
    cv2.imencode('.jpg', img_cropped, [int(cv2.IMWRITE_JPEG_QUALITY), 94])[1].tofile(cropped_path)

    return {
        "name": name,
        "id_card": id_card,
        "nation": nation,
        "temp_id_card_img": cropped_path
    }


def _ocr_result_text(line):
    """Read text from the RapidOCR result shape used by supported versions."""
    if not line or len(line) < 2:
        return ""
    value = line[1]
    if isinstance(value, (tuple, list)):
        value = value[0] if value else ""
    return str(value or "").strip()


def _ocr_result_confidence(line):
    if not line:
        return 0.0
    try:
        if len(line) > 2:
            return float(line[2])
        if len(line) > 1 and isinstance(line[1], (tuple, list)) and len(line[1]) > 1:
            return float(line[1][1])
    except (TypeError, ValueError):
        pass
    return 0.0


def _special_work_ocr_lines(ocr_results):
    lines = []
    for line in ocr_results or []:
        text = _ocr_result_text(line)
        if not text:
            continue
        x, y = 0.0, 0.0
        try:
            box = np.asarray(line[0], dtype=float)
            x = float(np.min(box[:, 0]))
            y = float(np.mean(box[:, 1]))
        except Exception:
            pass
        lines.append({"text": text, "compact": re.sub(r"\s+", "", text), "x": x, "y": y})
    return sorted(lines, key=lambda item: (round(item["y"] / 12.0), item["x"]))


def _special_work_ocr_score(ocr_results):
    score = 0.0
    for line in ocr_results or []:
        text = _ocr_result_text(line)
        if text:
            score += len(text) * (0.5 + max(0.0, min(1.0, _ocr_result_confidence(line))))
    return score


def _normalize_special_work_date(value):
    value = str(value or "")
    patterns = (
        r"(?<!\d)(20\d{2})\s*[年./\-]\s*(0?[1-9]|1[0-2])\s*[月./\-]\s*(0?[1-9]|[12]\d|3[01])\s*日?",
        r"(?<!\d)(20\d{2})(0[1-9]|1[0-2])([0-3]\d)(?!\d)",
    )
    for pattern in patterns:
        match = re.search(pattern, value)
        if not match:
            continue
        try:
            parsed = datetime.date(int(match.group(1)), int(match.group(2)), int(match.group(3)))
            return parsed.strftime("%Y%m%d")
        except ValueError:
            continue
    return ""


def _find_date_near_labels(lines, labels):
    for index, line in enumerate(lines):
        compact = line["compact"]
        if not any(label in compact for label in labels):
            continue
        for nearby in lines[index:index + 3]:
            matched_date = _normalize_special_work_date(nearby["compact"])
            if matched_date:
                return matched_date
    return ""


def _person_name_from_value(value):
    value = re.sub(r"^(姓名|持证人|姓\s*名)\s*[:：]?", "", str(value or ""))
    value = re.split(r"性别|民族|身份证|证件|出生|住址|工种", value)[0]
    cleaned = re.sub(r"[^\u4e00-\u9fff]", "", value)
    invalid = ("姓名", "持证人", "特种", "作业", "操作", "证书")
    return cleaned if 2 <= len(cleaned) <= 8 and not any(word in cleaned for word in invalid) else ""


def _value_after_labels(lines, labels, extractor=None):
    for index, line in enumerate(lines):
        compact = line["compact"]
        for label in labels:
            position = compact.find(label)
            if position < 0:
                continue
            candidates = [compact[position + len(label):].lstrip(":：-—_")]
            candidates.extend(next_line["compact"] for next_line in lines[index + 1:index + 3])
            for candidate in candidates:
                if not candidate or any(other in candidate for other in labels):
                    continue
                value = extractor(candidate) if extractor else candidate
                if value:
                    return value
    return ""


def _special_work_id_card(text):
    compact = re.sub(r"\s+", "", str(text or "")).upper()
    match = re.search(r"(?<!\d)([1-9]\d{5}[12]\d{3}[01]\d[0-3]\d{4}[\dX])(?![\dX])", compact)
    if match:
        return match.group(1)
    corrected = compact.replace("O", "0").replace("I", "1").replace("L", "1").replace("Z", "2").replace("S", "5")
    match = re.search(r"(?<!\d)([1-9]\d{5}[12]\d{3}[01]\d[0-3]\d{4}[\dX])(?![\dX])", corrected)
    return match.group(1) if match else ""


def _clean_certificate_name(value):
    value = re.sub(r"^(作业类别|操作项目|作业项目|证件名称|作业工种|准操项目|操作类别|作业种类)\s*[:：]?", "", str(value or ""))
    value = re.split(r"证书编号|证件编号|有效期|发证机关|发证单位|姓名|身份证", value)[0]
    value = re.sub(r"^[^\u4e00-\u9fffA-Za-z0-9]+|[^\u4e00-\u9fffA-Za-z0-9（）()\-]+$", "", value)
    return value[:60] if 2 <= len(value) <= 60 else ""


def _clean_issuing_authority(value):
    value = re.sub(r"^(发证机关|发证单位|发证部门|发证机构|签发机关)\s*[:：]?", "", str(value or ""))
    value = re.split(r"有效期|证书编号|证件编号|姓名|身份证", value)[0]
    value = re.sub(r"^[^\u4e00-\u9fffA-Za-z0-9]+|[^\u4e00-\u9fffA-Za-z0-9（）()\-]+$", "", value)
    return value[:80] if 2 <= len(value) <= 80 else ""


def extract_special_work_certificate_info(ocr_results):
    """Extract the fields needed by the special-work qualification register."""
    lines = _special_work_ocr_lines(ocr_results)
    full_text = "".join(line["compact"] for line in lines)
    name = _value_after_labels(lines, ("姓名", "持证人"), _person_name_from_value)
    id_card = _special_work_id_card(full_text)
    certificate_name = _value_after_labels(
        lines,
        ("作业类别", "操作项目", "作业项目", "证件名称", "作业工种", "准操项目", "操作类别", "作业种类"),
        _clean_certificate_name
    )
    if not certificate_name:
        work_keywords = ("电工", "焊", "司机", "信号", "高处", "登高", "制冷", "起重", "叉车", "有限空间", "架子", "压力")
        for line in lines:
            candidate = _clean_certificate_name(line["compact"])
            if candidate and any(keyword in candidate for keyword in work_keywords) and candidate not in ("特种作业操作证", "中华人民共和国特种作业操作证"):
                certificate_name = candidate
                break
    start_date = _find_date_near_labels(lines, ("起始日期", "初始日期", "初次取证", "初领日期", "发证日期", "生效日期"))
    end_date = _find_date_near_labels(lines, ("有效期至", "有效期限", "有效日期", "有效期"))
    all_dates = []
    for line in lines:
        value = _normalize_special_work_date(line["compact"])
        if value and value not in all_dates:
            all_dates.append(value)
    if not start_date and all_dates:
        start_date = all_dates[0]
    if not end_date and len(all_dates) > 1:
        end_date = all_dates[-1]
    issuing_authority = _value_after_labels(
        lines,
        ("发证机关", "发证单位", "发证部门", "发证机构", "签发机关"),
        _clean_issuing_authority
    )
    if not issuing_authority:
        issuer_keywords = ("应急管理", "市场监督", "安全生产", "管理局", "管理厅", "人民政府", "委员会", "住建")
        for line in lines:
            candidate = _clean_issuing_authority(line["compact"])
            if candidate and any(keyword in candidate for keyword in issuer_keywords):
                issuing_authority = candidate
                break
    certificate_number = _value_after_labels(
        lines,
        ("证书编号", "证件编号", "证书号"),
        lambda value: re.search(r"[A-Za-z0-9\-]{6,}", value).group(0) if re.search(r"[A-Za-z0-9\-]{6,}", value) else ""
    )
    return {
        "name": name,
        "id_card": id_card,
        "certificate_name": certificate_name,
        "start_date": start_date,
        "end_date": end_date,
        "issuing_authority": issuing_authority,
        "certificate_number": certificate_number,
        "raw_text": "\n".join(line["text"] for line in lines)[:2000],
        "recognized": bool(name or id_card or certificate_name or start_date or end_date or issuing_authority),
    }


def ocr_special_work_certificate_process(image_path):
    """OCR a special-work certificate and return register-ready fields."""
    engine = init_ppocrv6()
    source = _read_and_auto_orient(image_path)
    if source is None:
        raise ValueError("无法解析特殊工种证件图片。")
    candidates = (
        ("original", source),
        ("rotate_cw", cv2.rotate(source, cv2.ROTATE_90_CLOCKWISE)),
        ("rotate_ccw", cv2.rotate(source, cv2.ROTATE_90_COUNTERCLOCKWISE)),
        ("rotate_180", cv2.rotate(source, cv2.ROTATE_180)),
    )
    best_results = None
    best_score = -1.0
    best_orientation = "original"
    for index, (orientation, candidate) in enumerate(candidates):
        try:
            results, _ = engine(candidate)
        except Exception:
            results = None
        score = _special_work_ocr_score(results)
        if score > best_score:
            best_results = results
            best_score = score
            best_orientation = orientation
        if index == 0 and score >= 100:
            break
    if best_score < 30:
        filtered = _reduce_screen_pattern(source)
        try:
            filtered_results, _ = engine(filtered)
        except Exception:
            filtered_results = None
        filtered_score = _special_work_ocr_score(filtered_results)
        if filtered_score > best_score:
            best_results = filtered_results
            best_score = filtered_score
            best_orientation = "original_screen_filter"
    result = extract_special_work_certificate_info(best_results)
    result["orientation"] = best_orientation
    result["ocr_score"] = round(best_score, 2)
    print(
        f"[SpecialWork-OCR] direction={best_orientation}, score={best_score:.1f}, "
        f"name={'yes' if result['name'] else 'no'}, id={'yes' if result['id_card'] else 'no'}"
    )
    return result


# ================= 4. Word 登记卡自动生成 =================
def _center_text_in_spaces(text, total_spaces):
    width = sum(2 if ord(c) > 127 else 1 for c in text)
    if width >= total_spaces:
        return text
    left_pad = (total_spaces - width) // 2
    right_pad = total_spaces - width - left_pad
    return (" " * left_pad) + text + (" " * right_pad)

def generate_record_card(record_data, id_card_img_path):
    """
    根据登记卡模板生成 word 登记卡
    record_data: dict, 必须包含 姓名, 性别, 年龄, 联系电话, 岗位, 常住地址
    id_card_img_path: 裁剪后的身份证路径
    """
    template_path = '登记卡.docx'
    
    # 自动备份/提取桌面的登记卡模板
    if not os.path.exists(template_path):
        desktop_dir = os.path.join(os.environ.get('USERPROFILE', 'C:/Users/Administrator'), 'Desktop')
        src_template = os.path.join(desktop_dir, '身份证', '登记卡.docx')
        if os.path.exists(src_template):
            shutil.copy2(src_template, template_path)
            print("[Info] 从桌面身份证文件夹中成功复制 [登记卡.docx] 模板。")
        else:
            raise FileNotFoundError("本地和桌面上未找到 [登记卡.docx] 模板文件，请确认放置！")
            
    # 生成随机的目标 word 路径
    card_filename = f"card_{record_data['姓名']}_{uuid.uuid4().hex}.docx"
    output_path = os.path.join(CARDS_DIR, card_filename).replace('\\', '/')
    
    doc = Document(template_path)
    
    # 优先解析传入的录入日期 (北京时间)，如无则默认为当前北京时间
    created_at = record_data.get('created_at')
    record_date = None
    if created_at:
        try:
            if ' ' in created_at:
                record_date = datetime.datetime.strptime(created_at.split(' ')[0], "%Y-%m-%d")
            else:
                record_date = datetime.datetime.strptime(created_at, "%Y-%m-%d")
        except Exception:
            pass
            
    if not record_date:
        record_date = datetime.datetime.utcnow() + datetime.timedelta(hours=8)
    
    replacements = {
        '[[NAME]]': record_data.get('工作单位', ''),
        '[[name]]': record_data.get('工作单位', ''),
        '[[COMPANY]]': record_data.get('工作单位', ''),
        '[[GENDER]]': record_data.get('性别', ''),
        '[[AGE]]': str(record_data.get('年龄', '')),
        '[[PHONE]]': str(record_data.get('联系电话', '')),
        '[[YEAR]]': str(record_date.year),
        '[[MONTH]]': str(record_date.month),
        '[[DAY]]': str(record_date.day)
    }
    
    replaced_img = False
    paragraphs = [p for p in doc.paragraphs] + [p for t in doc.tables for r in t.rows for c in r.cells for p in c.paragraphs]
    
    def merge_similar_runs(paragraph):
        if len(paragraph.runs) <= 1:
            return
        i = 0
        while i < len(paragraph.runs) - 1:
            r1 = paragraph.runs[i]
            r2 = paragraph.runs[i+1]
            if (r1.underline == r2.underline and 
                r1.bold == r2.bold and 
                r1.italic == r2.italic and 
                r1.font.name == r2.font.name and 
                r1.font.size == r2.font.size):
                r1.text = r1.text + r2.text
                p_element = paragraph._element
                p_element.remove(r2._element)
                continue
            i += 1

    for p in paragraphs:
        merge_similar_runs(p)
            
        text = p.text.strip()
        if text == "工种":
            p.clear()
            r1 = p.add_run("工种")
            r1.font.name = "宋体"
            job_val = record_data.get("岗位", "普工")
            r2 = p.add_run(_center_text_in_spaces(job_val, 41))
            r2.font.name = "宋体"
            r2.underline = True
            continue
        elif text == "现住址":
            p.clear()
            r1 = p.add_run("现住址")
            r1.font.name = "宋体"
            addr_val = record_data.get("常住地址", "四川华庭生活区")
            r2 = p.add_run(_center_text_in_spaces(addr_val, 37))
            r2.font.name = "宋体"
            r2.underline = True
            continue
            
        if '[[ANCHOR_ID_CARD_HERE]]' in p.text and not replaced_img and id_card_img_path and os.path.exists(id_card_img_path):
            p.clear()
            run = p.add_run()
            # 尺寸：3.37" * 2.13" 身份证金牌尺寸
            run.add_picture(id_card_img_path, width=Inches(3.37), height=Inches(2.13))
            try:
                p.alignment = 1  # 居中对齐
            except:
                pass
            replaced_img = True
            
        for k, v in replacements.items():
            if k in p.text:
                for run in p.runs:
                    if k in run.text:
                        run.text = run.text.replace(k, v)
                        
    doc.save(output_path)
    return output_path

# ================= 5. 定期清理守护线程 =================
def cleanup_old_words():
    now = time.time()
    one_month_ago = now - 30 * 24 * 3600  # 一个月前 (30天)
    
    import glob
    
    # 1. 扫描 cards 目录下的 docx，并更新数据库
    if os.path.exists(CARDS_DIR):
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        for docx_file in glob.glob(os.path.join(CARDS_DIR, "*.docx")):
            try:
                mtime = os.path.getmtime(docx_file)
                if mtime < one_month_ago:
                    os.remove(docx_file)
                    filename = os.path.basename(docx_file)
                    cursor.execute("UPDATE records SET word_path = NULL WHERE word_path LIKE ?", (f"%{filename}",))
                    print(f"[Cleanup] 已成功自动清理 30 天前的登记卡 Word 文件: {filename}")
            except Exception as e:
                print(f"[Error] 自动清理文件 {docx_file} 失败: {e}")
        conn.commit()
        conn.close()
        
    # 2. uploads/idcards 中的是已入库的正式身份证裁切图，永久保留；
    #    仅在管理员主动清理对应录入资料时才会删除。

    # 3. 扫描 temp_ids 目录下的临时身份证照片，删除超过 1 天的
    if os.path.exists(TEMP_IDS_DIR):
        one_day_ago = now - 24 * 3600
        for temp_img in glob.glob(os.path.join(TEMP_IDS_DIR, "*")):
            try:
                mtime = os.path.getmtime(temp_img)
                if mtime < one_day_ago:
                    os.remove(temp_img)
                    print(f"[Cleanup] 已成功自动清理超过 1 天的临时身份证照片: {os.path.basename(temp_img)}")
            except Exception as e:
                print(f"[Error] 自动清理临时图片 {temp_img} 失败: {e}")

def start_cleanup_thread():
    def loop():
        while True:
            try:
                cleanup_old_words()
            except Exception as e:
                print(f"[Error] Cleanup thread exception: {e}")
            time.sleep(12 * 3600)  # 每 12 小时检查一次
            
    import threading
    t = threading.Thread(target=loop, daemon=True)
    t.start()
    print("[Info] 过期登记卡清理守护线程已成功开启（周期: 12小时）。")
