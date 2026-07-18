import os
import sys
import logging
from logging.handlers import RotatingFileHandler

# 自动双向重定向 stdout 和 stderr 到 app_stdout.log 中，保留控制台输出的同时记录到文件
# 日志按 10MB 轮转、最多保留 5 个历史文件，避免长期运行无限增长（N6）
class DualLogger:
    def __init__(self, filepath):
        self.terminal = sys.stdout
        self.handler = RotatingFileHandler(
            filepath, maxBytes=10 * 1024 * 1024, backupCount=5, encoding="utf-8"
        )
        self.handler.setFormatter(logging.Formatter("%(message)s"))
    def write(self, message):
        self.terminal.write(message)
        self.handler.stream.write(message)
    def flush(self):
        self.terminal.flush()
        self.handler.flush()

sys.stdout = DualLogger("app_stdout.log")
sys.stderr = DualLogger("app_stdout.log")

import os
import io
import re
import json
import sqlite3
import shutil
import hmac
import hashlib
import uuid
import time
import zipfile
import urllib.parse
from contextlib import contextmanager
from datetime import datetime, timedelta
from fastapi import FastAPI, HTTPException, UploadFile, File, Form, Header, Depends, Response, Request
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
import openpyxl
import csv

# 引入 OCR 核心
try:
    from ocr_handler import ocr_idcard_process, generate_record_card, start_cleanup_thread
    OCR_AVAILABLE = True
except Exception as e:
    OCR_AVAILABLE = False
    OCR_ERROR_MSG = str(e)
    print(f"[Warning] Failed to import ocr_handler, OCR features will be disabled: {e}")
    
    def ocr_idcard_process(*args, **kwargs):
        raise HTTPException(status_code=500, detail=f"OCR 引擎依赖缺失，请在服务器安装所需依赖（如 rapidocr_onnxruntime, modelscope, opencv-python-headless 等）。错误原因: {OCR_ERROR_MSG}")
        
    def generate_record_card(*args, **kwargs):
        raise HTTPException(status_code=500, detail=f"登记卡生成功能依赖缺失，请在服务器安装 python-docx。错误原因: {OCR_ERROR_MSG}")
        
    def start_cleanup_thread(*args, **kwargs):
        pass


DB_PATH = 'peixun.db'
UPLOAD_DIR = 'uploads'
TEMP_IDS_DIR = os.path.join(UPLOAD_DIR, 'temp_ids')

# 默认考试科目（供首次初始化及回退使用）。R3：原本在 5 处重复，现统一为常量。
DEFAULT_EXAM_SUBJECTS = [
    {"name": "普工", "file": "普工试题(1).xlsx"},
    {"name": "焊工", "file": "焊工题库(1).xlsx"},
    {"name": "探伤", "file": "探伤.xlsx"},
    {"name": "高处作业", "file": "高处作业(1).xlsx"},
    {"name": "吊装作业", "file": "吊装作业.xlsx"},
    {"name": "电工", "file": "电工题库.xlsx"},
    {"name": "叉车", "file": "叉车工.xlsx"}
]

def get_exam_file_map() -> dict:
    """读取 configs 中的考试科目配置，返回 {科目名: 文件名} 映射；缺失或损坏时回退默认。"""
    val = get_config('exam_subjects', '')
    if val:
        try:
            subjects = json.loads(val)
            return {s['name']: s['file'] for s in subjects}
        except Exception:
            pass
    return {s['name']: s['file'] for s in DEFAULT_EXAM_SUBJECTS}

def get_exam_subjects_list() -> list:
    """读取 configs 中的考试科目完整列表；缺失或损坏时回退默认。"""
    val = get_config('exam_subjects', '')
    if val:
        try:
            subjects = json.loads(val)
            if isinstance(subjects, list):
                return subjects
        except Exception:
            pass
    return [dict(s) for s in DEFAULT_EXAM_SUBJECTS]

def save_exam_subjects_list(subjects: list):
    """把考试科目列表写回 configs。"""
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("INSERT OR REPLACE INTO configs (key, value) VALUES ('exam_subjects', ?)", (json.dumps(subjects, ensure_ascii=False),))
        conn.commit()

# 考试题库内存缓存，结构: { exam_type: { "mtime": 12345, "answers": { ... } } }
EXAM_QUESTIONS_CACHE = {}

def get_exam_questions_answers(exam_type: str) -> dict:
    exam_file_map = get_exam_file_map()
    file_name = exam_file_map.get(exam_type)
    if not file_name:
        raise HTTPException(status_code=400, detail="未知的考试类型")
        
    file_path = os.path.join('shiti', file_name)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="题库文件不存在")
        
    mtime = os.path.getmtime(file_path)
    
    if exam_type in EXAM_QUESTIONS_CACHE:
        cache_data = EXAM_QUESTIONS_CACHE[exam_type]
        if cache_data.get('mtime') == mtime:
            return cache_data.get('answers')
            
    wb = openpyxl.load_workbook(file_path, read_only=True) # 使用 read_only 加速解析
    ws = wb.active
    shiti_answers = {}
    for row in ws.iter_rows(min_row=2, values_only=True):
        if len(row) >= 4 and row[1] and row[3]:
            question_text = str(row[1]).strip()
            ans = str(row[3]).strip()
            if ans in ['对', '正确', '√', 'T', 'True']:
                ans = '对'
            elif ans in ['错', '错误', '×', 'F', 'False']:
                ans = '错'
            shiti_answers[question_text] = ans
            
    EXAM_QUESTIONS_CACHE[exam_type] = {
        'mtime': mtime,
        'answers': shiti_answers
    }
    return shiti_answers

# S4: Token 签名密钥优先从环境变量读取，避免硬编码进源码造成泄露后可伪造任意 Token。
# 未设置环境变量时使用一个进程级随机值（重启后所有已签发 Token 失效），保证默认安全。
SECRET_KEY = os.environ.get("PEIXUN_SECRET_KEY", "").encode("utf-8")
if not SECRET_KEY:
    SECRET_KEY = os.urandom(32)
    print("[Warning] 未设置环境变量 PEIXUN_SECRET_KEY，已使用随机密钥（重启后所有登录将失效）。")

def encrypt_pwd(raw_pwd: str) -> str:
    salt = uuid.uuid4().hex
    hashed = hashlib.sha256((raw_pwd + salt).encode('utf-8')).hexdigest()
    return f"{salt}${hashed}"

def verify_pwd(raw_pwd: str, stored_pwd_str: str) -> bool:
    try:
        salt, hashed = stored_pwd_str.split("$")
        check_hashed = hashlib.sha256((raw_pwd + salt).encode('utf-8')).hexdigest()
        return hmac.compare_digest(check_hashed, hashed)
    except ValueError:
        # L8: 历史明文密码兼容分支。注册接口已统一使用 encrypt_pwd 加盐哈希，
        # 此分支仅兼容迁移前的明文记录——登录成功后 login 接口会自动升级为哈希。
        # 待确认数据库无明文密码后可移除此分支。
        return hmac.compare_digest(raw_pwd, stored_pwd_str)

def generate_token(user_id: int, role: str, username: str, expire_seconds: int) -> str:
    """签发 HMAC 签名 token。

    token 结构: user_id:role:username:issued_at:expire_at:signature
    issued_at / expire_at 均为 unix 秒级时间戳。expire_at 纳入签名，
    使单条 token 的有效期不可被篡改。
    """
    issued_at = int(time.time())
    expire_at = issued_at + int(expire_seconds)
    payload = f"{user_id}:{role}:{username}:{issued_at}:{expire_at}"
    signature = hmac.new(SECRET_KEY, payload.encode('utf-8'), hashlib.sha256).hexdigest()
    return f"{payload}:{signature}"

def verify_token(token: str) -> dict:
    try:
        parts = token.split(":")
        if len(parts) != 6:
            return None
        user_id, role, username, issued_at, expire_at, signature = parts
        payload = f"{user_id}:{role}:{username}:{issued_at}:{expire_at}"
        expected_signature = hmac.new(SECRET_KEY, payload.encode('utf-8'), hashlib.sha256).hexdigest()
        if not hmac.compare_digest(expected_signature, signature):
            return None
        # 以 token 自带的过期时间为准（登录时按 remember 勾选决定）
        if int(time.time()) > int(expire_at):
            return None
        return {
            "id": int(user_id),
            "role": role,
            "username": username
        }
    except Exception:
        return None

def get_config(key: str, default: str) -> str:
    try:
        with get_db() as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT value FROM configs WHERE key = ?", (key,))
            row = cursor.fetchone()
            return row[0] if row else default
    except Exception:
        return default

def beijing_now() -> datetime:
    """统一返回北京时间 (UTC+8) 当前时间，不受服务器本地时区影响（L4）。"""
    return datetime.utcnow() + timedelta(hours=8)

def is_exam_open() -> bool:
    # 强制以北京时间 (UTC+8) 校验，不受服务器本地时区设置影响
    now_dt = beijing_now()
    now_time = now_dt.time()
    
    start_str = get_config('exam_start_time', '08:00:00')
    end_str = get_config('exam_end_time', '12:00:00')
    
    try:
        start_time = datetime.strptime(start_str, "%H:%M:%S").time()
        end_time = datetime.strptime(end_str, "%H:%M:%S").time()
    except Exception:
        start_time = datetime.strptime("08:00:00", "%H:%M:%S").time()
        end_time = datetime.strptime("12:00:00", "%H:%M:%S").time()
        
    return start_time <= now_time <= end_time

# R5: 统一的数据库连接上下文管理器，自动关闭连接。
@contextmanager
def get_db():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    try:
        yield conn
    finally:
        conn.close()

# S5: 标准 18 位身份证号校验（地址码 6 位 + 8 位生日合法 + 3 位顺序码 + 1 位 ISO 7064 Mod 11 校验码）。
# 替换原先错误且无依据的“第 11 位必须为 0 或 1”判断（该位是出生日首位，合法范围 0-3）。
_ID_CARD_WEIGHTS = (7, 9, 10, 5, 8, 4, 2, 1, 6, 3, 7, 9, 10, 5, 8, 4, 2)
_ID_CARD_CHECK = ('1', '0', 'X', '9', '8', '7', '6', '5', '4', '3', '2')

def validate_id_card(id_card: str) -> bool:
    """校验身份证号合法性，合法返回 True。"""
    if not id_card or len(id_card) != 18:
        return False
    if not id_card[:17].isdigit():
        return False
    # 校验生日合法性（支持 1800-2099 年）
    try:
        birthday = datetime.strptime(id_card[6:14], "%Y%m%d")
        if birthday.year < 1800 or birthday.year > 2099:
            return False
    except ValueError:
        return False
    total = sum(int(id_card[i]) * _ID_CARD_WEIGHTS[i] for i in range(17))
    return _ID_CARD_CHECK[total % 11] == id_card[17].upper()

# S2: 校验前端回传的身份证裁剪图路径确实位于 uploads/temp_ids 目录内，
# 防止攻击者传任意服务器路径（如 peixun.db / 登记卡.docx）被复制并经 /uploads 公开下载。
def safe_temp_id_path(path: str) -> bool:
    if not path:
        return False
    base = os.path.realpath(TEMP_IDS_DIR)
    real = os.path.realpath(path)
    # 必须在 temp_ids 目录之下，且文件真实存在
    return real.startswith(base + os.sep) and os.path.isfile(real)

# R4: 过滤字符串中不能用于文件名的字符（跨平台），用于导出文件命名。
_FILENAME_INVALID_CHARS = r'\/:*?"<>|'

def safe_filename_part(s: str) -> str:
    if not s:
        return ""
    return "".join(c for c in str(s) if c not in _FILENAME_INVALID_CHARS)

# 门禁/恢复导入表 CSV 的说明注释（无法读取 06.03.csv 模板时回退使用）。
_GATE_CSV_DEFAULT_COMMENTS = [
    "・ *号为必填项；,,,,,,,,,\n",
    "・ 填写数字编码代替属性值；,,,,,,,,,\n",
    "・ 使用EXCEL编辑导入文件时，请将单元格的格式修改为文本格式，避免数字文本自动转换为科学计数文本；,,,,,,,,,\n",
    ",,,,,,,,,\n",
    "1、姓名*：1～32个字符；不能包含 ' / \\: * ? \" < > | 这些特殊字符；,,,,,,,,,\n",
    "2、性别*：1（男）、2（女）、0（未知）；,,,,,,,,,\n",
    "3、组织路径*：填写从选择导入的组织名称开始，至目标组织的完整名称路径；,,,,,,,,,\n",
    "4、证件类型*：111（身份证）、414（护照）、113（户口簿）、335（驾驶证）、131（工作证）、133（学生证）、114（军官证）、990（其他）；,,,,,,,,,\n",
    "5、证件号码*：1~20个字符；只允许输入数字和字母；,,,,,,,,,\n",
    "6、工号：1~32个字符；只允许输入数字、字母和汉字；,,,,,,,,,\n",
    "7、手机号码：1-20位数字；,,,,,,,,,\n",
    "8、拼音：人员姓名拼音；,,,,,,,,,\n",
    "9、所属区域：0 ～128个字符；不能包含 ' / \\ : * ? \" < > | 这些特殊字符；,,,,,,,,,\n",
    "10、卡号：8~20个字符；只允许输入数字和大写字母。,,,,,,,,,\n"
]

def _load_gate_csv_comments() -> list:
    """优先读取 06.03.csv 模板首部说明注释，失败则使用默认注释。"""
    template_path = '06.03.csv'
    if os.path.exists(template_path):
        try:
            with open(template_path, 'r', encoding='gbk') as f:
                lines = f.readlines()
                for idx, line in enumerate(lines):
                    if '*姓名' in line:
                        return lines[:idx]
        except Exception:
            pass
    return list(_GATE_CSV_DEFAULT_COMMENTS)

def build_gate_csv(records) -> bytes:
    """根据 records（sqlite Row 列表，需含 name/gender/company/region_auth/id_card/phone）
    生成 GBK 编码的门禁导入表 CSV 字节，供下载接口复用。"""
    csv_output = io.StringIO()
    csv_output.write("".join(_load_gate_csv_comments()))
    writer = csv.writer(csv_output)
    writer.writerow(['*姓名', '*性别', '*组织路径', '*证件类型', '*证件号码', '工号', '手机号码', '拼音', '所属区域', '卡号'])
    for r in records:
        gender_code = '0'
        if r['gender'] == '男':
            gender_code = '1'
        elif r['gender'] == '女':
            gender_code = '2'
        company = r['company'] if r['company'] else ""
        region = r['region_auth'] if r['region_auth'] else ""
        org_path = f"{company}/{region}" if region else company
        writer.writerow([r['name'], gender_code, org_path, '111', r['id_card'], '', r['phone'], '', region, ''])
    csv_data = csv_output.getvalue().encode('gbk', errors='ignore')
    csv_output.close()
    return csv_data

def pack_photos_zip(records) -> bytes:
    """根据 records（需含 photo_path/name/id_card）生成照片压缩包字节，自动处理文件名净化与重名。"""
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        added_filenames = set()
        for r in records:
            photo_path = r['photo_path']
            if not photo_path or not os.path.exists(photo_path):
                continue
            file_ext = os.path.splitext(photo_path)[1] or '.jpg'
            base_filename = f"{safe_filename_part(r['name'])}_{safe_filename_part(r['id_card'])}"
            filename = f"{base_filename}{file_ext}"
            counter = 1
            while filename in added_filenames:
                filename = f"{base_filename}_{counter}{file_ext}"
                counter += 1
            added_filenames.add(filename)
            zip_file.write(photo_path, arcname=filename)
    return zip_buffer.getvalue()

def _csv_download_response(content: bytes, filename: str):
    """生成带 UTF-8 文件名的 CSV 下载响应。"""
    return Response(
        content=content,
        media_type="text/csv",
        headers={"Content-Disposition": f"attachment; filename*=utf-8''{urllib.parse.quote(filename)}"}
    )

def _zip_download_response(content: bytes, filename: str):
    """生成带 UTF-8 文件名的 ZIP 下载响应。"""
    return Response(
        content=content,
        media_type="application/zip",
        headers={"Content-Disposition": f"attachment; filename*=utf-8''{urllib.parse.quote(filename)}"}
    )

def parse_id_list(ids: str) -> list:
    """把 "1,2,3" 解析为 [1,2,3]，校验非空。"""
    if not ids:
        raise HTTPException(status_code=400, detail="请选择需要下载的记录")
    id_list = [int(x) for x in ids.split(',') if x.strip()]
    if not id_list:
        raise HTTPException(status_code=400, detail="没有勾选任何有效的人员记录")
    return id_list

def fetch_records_by_ids(cursor, id_list: list, gate_only: bool):
    """按 id 列表查询 records（含 company），gate_only=True 时仅查门禁恢复待处理记录。
    返回 (records, placeholders)。"""
    placeholders = ','.join(['?'] * len(id_list))
    where = f"r.id IN ({placeholders})"
    if gate_only:
        where += " AND r.gate_restore_status = 'pending'"
    cursor.execute(f'''
    SELECT r.*
    FROM records r
    WHERE {where}
    ORDER BY r.is_gate_downloaded ASC, r.created_at DESC
    ''', id_list)
    return cursor.fetchall(), placeholders


def get_latest_training_record(cursor, name: str, id_last6: str):
    """Return the newest training record for a verified person.

    ``records.company`` is captured at training time.  The fallback keeps
    records created before this field was introduced usable after migration.
    """
    cursor.execute('''
        SELECT r.*, COALESCE(NULLIF(r.company, ''), u.company, '') AS training_company
        FROM records r
        LEFT JOIN users u ON r.user_id = u.id
        WHERE r.name = ? AND SUBSTR(r.id_card, -6) = ?
        ORDER BY r.created_at DESC, r.id DESC
        LIMIT 1
    ''', (name, id_last6))
    return cursor.fetchone()

# 初始化数据库
def init_db():
    if not os.path.exists(UPLOAD_DIR):
        os.makedirs(UPLOAD_DIR)
        
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    # 开启 WAL 模式以提升并发读写性能
    cursor.execute("PRAGMA journal_mode=WAL;")
    
    # 用户表
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS users (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        username TEXT UNIQUE NOT NULL, -- 登录手机号/用户名
        password TEXT NOT NULL,
        real_name TEXT NOT NULL,
        company TEXT NOT NULL,
        status TEXT NOT NULL DEFAULT 'pending', -- pending, approved, rejected
        role TEXT NOT NULL DEFAULT 'user' -- user, admin
    )
    ''')
    
    # 录入信息表
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS records (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER,
        photo_path TEXT,
        name TEXT,
        nation TEXT,
        id_card TEXT,
        phone TEXT,
        address TEXT,
        job TEXT,
        education TEXT,
        region_auth TEXT,
        gender TEXT,
        age INTEGER,
        company TEXT DEFAULT '',
        remark TEXT DEFAULT '',
        created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY(user_id) REFERENCES users(id)
    )
    ''')
    
    # 默认管理员
    cursor.execute("SELECT * FROM users WHERE username = 'admin'")
    if not cursor.fetchone():
        cursor.execute(
            "INSERT INTO users (username, password, real_name, company, status, role) VALUES (?, ?, ?, ?, ?, ?)",
            ('admin', encrypt_pwd('admin123'), '系统管理员', '管理部', 'approved', 'admin')
        )
        
    cursor.execute("SELECT * FROM users WHERE username = 'admin2'")
    if not cursor.fetchone():
        cursor.execute(
            "INSERT INTO users (username, password, real_name, company, status, role) VALUES (?, ?, ?, ?, ?, ?)",
            ('admin2', encrypt_pwd('admin1234'), '普通管理员', '管理部', 'approved', 'admin')
        )
        
    # 检测并为 records 表添加 is_gate_downloaded 列自愈逻辑
    try:
        cursor.execute("ALTER TABLE records ADD COLUMN is_gate_downloaded INTEGER DEFAULT 0")
        conn.commit()
    except sqlite3.OperationalError:
        pass

    # 检测并为 records 表添加 gate_restore_status 列自愈逻辑
    try:
        cursor.execute("ALTER TABLE records ADD COLUMN gate_restore_status TEXT DEFAULT NULL")
        conn.commit()
    except sqlite3.OperationalError:
        pass

    # 检测并为 records 表添加 is_restore_downloaded 列自愈逻辑
    try:
        cursor.execute("ALTER TABLE records ADD COLUMN is_restore_downloaded INTEGER DEFAULT 0")
        conn.commit()
    except sqlite3.OperationalError:
        pass

    # 检测并为 records 表添加 word_path 列自愈逻辑
    try:
        cursor.execute("ALTER TABLE records ADD COLUMN word_path TEXT DEFAULT NULL")
        conn.commit()
    except sqlite3.OperationalError:
        pass

    # Preserve the company that applied to each training entry.  It must not
    # change when the recorder later moves to another company.
    try:
        cursor.execute("ALTER TABLE records ADD COLUMN company TEXT DEFAULT ''")
        conn.commit()
    except sqlite3.OperationalError:
        pass
    try:
        cursor.execute("ALTER TABLE records ADD COLUMN remark TEXT DEFAULT ''")
        conn.commit()
    except sqlite3.OperationalError:
        pass
    cursor.execute('''
        UPDATE records
        SET company = COALESCE(
            NULLIF(company, ''),
            (SELECT users.company FROM users WHERE users.id = records.user_id),
            ''
        )
        WHERE company IS NULL OR TRIM(company) = ''
    ''')
    
    # 答题记录表
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS exam_records (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        name TEXT NOT NULL,
        company TEXT NOT NULL,
        exam_type TEXT NOT NULL,
        score INTEGER NOT NULL,
        answered_count INTEGER NOT NULL,
        correct_count INTEGER NOT NULL,
        duration TEXT NOT NULL,
        created_at DATETIME DEFAULT CURRENT_TIMESTAMP
    )
    ''')
    
    # 答题详情表
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS exam_details (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        exam_record_id INTEGER NOT NULL,
        question TEXT NOT NULL,
        user_answer TEXT NOT NULL,
        correct_answer TEXT NOT NULL,
        is_correct INTEGER NOT NULL DEFAULT 0,
        FOREIGN KEY(exam_record_id) REFERENCES exam_records(id)
    )
    ''')
    # 系统配置表
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS configs (
        key TEXT PRIMARY KEY,
        value TEXT NOT NULL
    )
    ''')
    cursor.execute("INSERT OR IGNORE INTO configs (key, value) VALUES ('exam_start_time', '08:00:00')")
    cursor.execute("INSERT OR IGNORE INTO configs (key, value) VALUES ('exam_end_time', '12:00:00')")
    cursor.execute("INSERT OR IGNORE INTO configs (key, value) VALUES ('regions', '三元肥,尿素塔')")
        
    # 答题审批表
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS exam_approvals (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        name TEXT NOT NULL,
        id_last6 TEXT NOT NULL,
        status TEXT NOT NULL DEFAULT 'pending',
        company TEXT DEFAULT NULL,
        created_at DATETIME DEFAULT CURRENT_TIMESTAMP
    )
    ''')
    cursor.execute("CREATE INDEX IF NOT EXISTS idx_exam_approvals_status ON exam_approvals (status)")

    # 建立索引以优化检索和排序性能
    cursor.execute("CREATE INDEX IF NOT EXISTS idx_records_name ON records (name)")
    cursor.execute("CREATE INDEX IF NOT EXISTS idx_records_created_at ON records (created_at)")
    cursor.execute("CREATE INDEX IF NOT EXISTS idx_records_company ON records (company)")
    cursor.execute("CREATE INDEX IF NOT EXISTS idx_records_id_card_created_at ON records (id_card, created_at)")
    cursor.execute("CREATE INDEX IF NOT EXISTS idx_exam_records_created_at ON exam_records (created_at)")
    cursor.execute("CREATE INDEX IF NOT EXISTS idx_exam_records_name ON exam_records (name)")
    cursor.execute("CREATE INDEX IF NOT EXISTS idx_exam_records_company ON exam_records (company)")
    cursor.execute("CREATE INDEX IF NOT EXISTS idx_users_company ON users (company)")
    
    # 自愈逻辑：给 users 增加 is_deleted 和 created_at 字段
    try:
        cursor.execute("ALTER TABLE users ADD COLUMN is_deleted INTEGER DEFAULT 0")
    except sqlite3.OperationalError:
        pass
    try:
        cursor.execute("ALTER TABLE users ADD COLUMN created_at DATETIME DEFAULT CURRENT_TIMESTAMP")
    except sqlite3.OperationalError:
        pass

    # 人员信息修改申请表
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS record_updates (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        record_id INTEGER NOT NULL,
        user_id INTEGER NOT NULL,
        name TEXT NOT NULL,
        nation TEXT NOT NULL,
        id_card TEXT NOT NULL,
        phone TEXT NOT NULL,
        address TEXT NOT NULL,
        job TEXT NOT NULL,
        education TEXT NOT NULL,
        region_auth TEXT,
        remark TEXT DEFAULT '',
        photo_path TEXT,
        id_card_img_path TEXT,
        status TEXT NOT NULL DEFAULT 'pending',
        created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY(record_id) REFERENCES records(id),
        FOREIGN KEY(user_id) REFERENCES users(id)
    )
    ''')
    cursor.execute("CREATE INDEX IF NOT EXISTS idx_record_updates_status ON record_updates (status)")
    try:
        cursor.execute("ALTER TABLE record_updates ADD COLUMN remark TEXT DEFAULT ''")
    except sqlite3.OperationalError:
        pass
    
    # 历史数据时区迁移自愈逻辑（将原本依赖 SQLite DEFAULT CURRENT_TIMESTAMP 存入的 UTC 时间迁移至北京时间 +8 小时，仅迁移一次）
    cursor.execute("SELECT value FROM configs WHERE key = 'timezone_fixed'")
    row = cursor.fetchone()
    if not row or row[0] != 'true':
        try:
            cursor.execute("UPDATE records SET created_at = datetime(created_at, '+8 hours') WHERE created_at IS NOT NULL")
            cursor.execute("UPDATE exam_records SET created_at = datetime(created_at, '+8 hours') WHERE created_at IS NOT NULL")
            cursor.execute("INSERT OR REPLACE INTO configs (key, value) VALUES ('timezone_fixed', 'true')")
            print("Successfully migrated historical records timezone from UTC to UTC+8")
        except Exception as e:
            print(f"Failed to migrate timezone: {e}")
            
    conn.commit()
    conn.close()

init_db()
# 预先加载并常驻 OCR 模型，防止首笔请求慢
if OCR_AVAILABLE:
    try:
        from ocr_handler import init_ppocrv6
        init_ppocrv6()
    except Exception as e:
        print(f"[Warning] Startup preloading OCR model failed: {e}")
    
    try:
        start_cleanup_thread()
    except Exception as e:
        print(f"[Warning] Failed to start cleanup thread: {e}")


app = FastAPI(title="培训信息录入系统")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,  # S5: 系统通过 Header 传 Token，无需 cookie 凭证；与 * 同用更安全
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.middleware("http")
async def log_requests(request, call_next):
    start_time = time.time()
    response = await call_next(request)
    duration = time.time() - start_time
    # S4: 脱敏 query_params 中的 token/authorization，避免凭证泄露到日志文件
    safe_qp = ""
    if request.query_params:
        safe_items = []
        for k, v in request.query_params.multi_items():
            if k.lower() in ("token", "authorization"):
                safe_items.append(f"{k}=***")
            else:
                safe_items.append(f"{k}={v}")
        safe_qp = "&".join(safe_items)
    print(f"DEBUG_REQ: {request.method} {request.url.path} {safe_qp} - Status: {response.status_code} - Duration: {duration:.2f}s")
    sys.stdout.flush()
    return response

# S3: 保护 uploads 下的敏感目录（身份证照片、登记卡 Word、临时身份证图），
# 任何人访问这些资源必须携带有效 token，防止身份证照片被任意下载
@app.middleware("http")
async def protect_uploads(request, call_next):
    path = request.url.path
    protected_prefixes = ("/uploads/idcards/", "/uploads/cards/", "/uploads/temp_ids/")
    if any(path.startswith(p) for p in protected_prefixes):
        token = request.query_params.get("token") or request.headers.get("authorization")
        if not token or not verify_token(token):
            return JSONResponse(status_code=401, content={"detail": "未授权访问该资源"})
    return await call_next(request)

# 身份证解析逻辑
def parse_id_card(id_card_num):
    if not id_card_num or len(id_card_num) != 18:
        return "未知", 0
    try:
        birth_year_str = id_card_num[6:10]
        birth_year = int(birth_year_str)
        # 获取当前年份
        current_year = datetime.now().year
        age = current_year - birth_year
        
        gender_digit = int(id_card_num[16])
        gender = "男" if gender_digit % 2 != 0 else "女"
        return gender, age
    except Exception:
        return "未知", 0

# 鉴权依赖
def get_current_user(authorization: str = Header(None)):
    if not authorization:
        raise HTTPException(status_code=401, detail="请重新登录")

    token_info = verify_token(authorization)
    if not token_info:
        raise HTTPException(status_code=401, detail="请重新登录")
    
    try:
        user_id = token_info["id"]
        username = token_info["username"]
        
        conn = sqlite3.connect(DB_PATH)
        conn.row_factory = sqlite3.Row
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM users WHERE id = ? AND username = ?", (user_id, username))
        user = cursor.fetchone()
        conn.close()
        
        if not user:
            raise HTTPException(status_code=401, detail="请重新登录")

        if user['status'] != 'approved' and user['role'] != 'admin':
            raise HTTPException(status_code=403, detail="账号尚未被审批通过，请耐心等待")

        return user
    except HTTPException:
        raise  # L1: 保留上面主动抛出的具体错误信息（如"账号尚未审批"），不被通用异常吞掉
    except Exception:
        raise HTTPException(status_code=401, detail="请重新登录")

def get_admin_user(current_user = Depends(get_current_user)):
    if current_user['role'] != 'admin':
        raise HTTPException(status_code=403, detail="无管理员权限")
    return current_user

# ---------------- API 接口 ----------------

@app.post("/api/save_ppt")
def save_ppt(html_content: str = Form(...), admin = Depends(get_admin_user)):
    # S1: 该接口可直接覆盖 static/ppt.html 内容，必须限定管理员，避免匿名篡改/XSS。
    try:
        static_dir = os.path.abspath("static")
        ppt_path = os.path.abspath(os.path.join(static_dir, "ppt.html"))
        with open(ppt_path, "w", encoding="utf-8") as f:
            f.write(html_content)
        return {"code": 200, "message": "保存成功！"}
    except Exception as e:
        print(f"[Error] save_ppt 失败: {e}")
        raise HTTPException(status_code=500, detail="保存失败")

@app.post("/api/register")
def register(username: str = Form(...), password: str = Form(...), real_name: str = Form(...), company: str = Form(...)):
    if not real_name.strip():
        raise HTTPException(status_code=400, detail="真实姓名不能为空")
    if not username.strip():
        raise HTTPException(status_code=400, detail="联系电话/用户名不能为空")
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    try:
        try:
            now_str = beijing_now().strftime("%Y-%m-%d %H:%M:%S")
            cursor.execute(
                "INSERT INTO users (username, password, real_name, company, status, role, created_at) VALUES (?, ?, ?, ?, 'pending', 'user', ?)",
                (username.strip(), encrypt_pwd(password.strip()), real_name.strip(), company.strip(), now_str)
            )
        except sqlite3.OperationalError as e:
            if "no column named created_at" in str(e):
                cursor.execute(
                    "INSERT INTO users (username, password, real_name, company, status, role) VALUES (?, ?, ?, ?, 'pending', 'user')",
                    (username.strip(), encrypt_pwd(password.strip()), real_name.strip(), company.strip())
                )
            else:
                raise e
        conn.commit()
        return {"code": 200, "message": "注册成功，请等待管理员审批！"}
    except sqlite3.IntegrityError:
        raise HTTPException(status_code=400, detail="该账号（电话）已存在，请直接登录")
    finally:
        conn.close()

@app.post("/api/login")
def login(request: Request, username: str = Form(...), password: str = Form(...), remember: str = Form("false")):
    client_ip = request.client.host if request.client else "unknown"
    ua = request.headers.get("user-agent", "unknown")
    
    # 清理输入的前后空格，防止用户复制时带入不可见字符，与注册行为保持一致
    username_clean = username.strip()
    password_clean = password.strip()
    
    # 打印原始输入，以便排查不同手机输入法自动加空格或大小写的问题
    print(f"[Login-Debug] 尝试登录 | 原始输入: username={repr(username)}, pwd_len={len(password)} | 清理后: username={repr(username_clean)}, pwd_len={len(password_clean)}")
    sys.stdout.flush()

    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM users WHERE username = ?", (username_clean,))
    user = cursor.fetchone()

    if not user:
        print(f"[Login-Debug] 登录失败: 用户不存在. 输入的用户名={repr(username_clean)} (IP={client_ip})")
        sys.stdout.flush()
        conn.close()
        raise HTTPException(status_code=400, detail="用户名或密码错误")

    if not verify_pwd(password_clean, user['password']):
        print(f"[Login-Debug] 登录失败: 密码不匹配. 用户名={repr(username_clean)} (IP={client_ip})")
        sys.stdout.flush()
        conn.close()
        raise HTTPException(status_code=400, detail="用户名或密码错误")

    # 平滑升级明文密码
    if '$' not in user['password']:
        try:
            cursor.execute("UPDATE users SET password = ? WHERE id = ?", (encrypt_pwd(password), user['id']))
            conn.commit()
        except Exception:
            pass

    conn.close()

    if user['status'] == 'pending' and user['role'] != 'admin':
        return {"code": 300, "message": "您的注册申请正在审批中，暂无法登录。"}

    if user['status'] == 'rejected':
        return {"code": 301, "message": "您的注册申请已被拒绝，请重新注册或联系管理员。"}

    # 勾选"保持登录"签发 30 天 token；不勾选签发 1 天 token
    # （前端在不勾选时使用 sessionStorage，关闭浏览器即清除，实现"关闭浏览器失效"）
    remember_flag = str(remember).strip().lower() in ("true", "1", "yes", "on")
    expire_seconds = 30 * 24 * 3600 if remember_flag else 1 * 24 * 3600
    token = generate_token(user['id'], user['role'], user['username'], expire_seconds)
    return {
        "code": 200,
        "message": "登录成功",
        "data": {
            "token": token,
            "role": user['role'],
            "username": user['username'],
            "real_name": user['real_name'],
            "company": user['company']
        }
    }

# 获取当前用户状态
@app.get("/api/user/status")
def user_status(current_user = Depends(get_current_user)):
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    # 额外去 records 表查询该用户最近提交的一条数据以用于回填（民族、常住地址、工种、学历、区域权限）
    cursor.execute('''
    SELECT nation, address, job, education, region_auth 
    FROM records 
    WHERE user_id = ? 
    ORDER BY id DESC LIMIT 1
    ''', (current_user['id'],))
    last_record = cursor.fetchone()
    conn.close()
    
    saved_data = {}
    if last_record:
        saved_data = {
            "nation": last_record["nation"],
            "address": last_record["address"],
            "job": last_record["job"],
            "education": last_record["education"],
            "region_auth": last_record["region_auth"]
        }
    else:
        # 如果是首次录入（没有历史记录），则为空
        saved_data = {
            "nation": "",
            "address": "",
            "job": "",
            "education": "",
            "region_auth": ""
        }

    return {
        "code": 200,
        "data": {
            "id": current_user["id"],
            "username": current_user["username"],
            "real_name": current_user["real_name"],
            "company": current_user["company"],
            "status": current_user["status"],
            "role": current_user["role"],
            "saved_fields": saved_data
        }
    }

# 获取待审批用户列表（仅管理员）
# 获取所有注册普通用户列表（包括 pending, approved, rejected，仅管理员）
@app.get("/api/admin/pending")
def get_pending_users(admin = Depends(get_admin_user)):
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    try:
        cursor.execute("SELECT id, username, real_name, company, status, created_at FROM users WHERE role != 'admin' AND is_deleted = 0 ORDER BY id DESC")
        users = [dict(row) for row in cursor.fetchall()]
    except sqlite3.OperationalError:
        try:
            cursor.execute("SELECT id, username, real_name, company, status FROM users WHERE role != 'admin' AND is_deleted = 0 ORDER BY id DESC")
            users = [dict(row) for row in cursor.fetchall()]
        except sqlite3.OperationalError:
            cursor.execute("SELECT id, username, real_name, company, status FROM users WHERE role != 'admin' ORDER BY id DESC")
            users = [dict(row) for row in cursor.fetchall()]
    conn.close()
    return {"code": 200, "data": users}

# 审批操作（仅管理员）
@app.post("/api/admin/approve")
def approve_user(user_id: int = Form(...), action: str = Form(...), admin = Depends(get_admin_user)):
    if action not in ['approved', 'rejected']:
        raise HTTPException(status_code=400, detail="非法审批操作")
        
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("UPDATE users SET status = ? WHERE id = ?", (action, user_id))
    conn.commit()
    conn.close()
    return {"code": 200, "message": "审批成功"}

# 删除注册用户（仅管理员，物理删除以防再度登录）
@app.post("/api/admin/user/delete")
def delete_user(user_id: int = Form(...), admin = Depends(get_admin_user)):
    if admin['username'] == 'admin2':
        raise HTTPException(status_code=403, detail="该账户无删除权限")
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    # 安全检查，防止删除管理员
    cursor.execute("SELECT role FROM users WHERE id = ?", (user_id,))
    user = cursor.fetchone()
    if not user:
        conn.close()
        raise HTTPException(status_code=404, detail="用户不存在")
    if user[0] == 'admin':
        conn.close()
        raise HTTPException(status_code=400, detail="无权删除管理员账户")
        
    cursor.execute("UPDATE users SET is_deleted = 1, username = username || '_deleted_' || id WHERE id = ?", (user_id,))
    conn.commit()
    conn.close()
    return {"code": 200, "message": "用户删除成功（已标记为已删除）"}

# 删除培训人员记录（仅管理员，物理删除关联照片）
@app.post("/api/admin/record/delete")
def delete_record(record_id: int = Form(...), admin = Depends(get_admin_user)):
    if admin['username'] == 'admin2':
        raise HTTPException(status_code=403, detail="该账户无删除权限")
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    # 1. 获取照片路径，以便随后进行物理删除
    cursor.execute("SELECT photo_path FROM records WHERE id = ?", (record_id,))
    row = cursor.fetchone()
    if not row:
        conn.close()
        raise HTTPException(status_code=404, detail="该人员记录不存在")
        
    photo_path = row[0]
    
    # 2. 从数据库删除记录
    cursor.execute("DELETE FROM records WHERE id = ?", (record_id,))
    conn.commit()
    conn.close()
    
    # 3. 物理删除照片文件
    if photo_path and os.path.exists(photo_path):
        try:
            os.remove(photo_path)
        except Exception:
            pass
            
    return {"code": 200, "message": "记录删除成功"}


# 管理员修改普通用户信息（仅管理员）
@app.post("/api/admin/user/update")
def admin_update_user(
    user_id: int = Form(...),
    username: str = Form(...),
    real_name: str = Form(...),
    company: str = Form(...),
    password: str = Form(None),
    admin = Depends(get_admin_user)
):
    username = username.strip()
    real_name = real_name.strip()
    company = company.strip()
    
    if not username:
        raise HTTPException(status_code=400, detail="申请账号（电话）不能为空")
    if not real_name:
        raise HTTPException(status_code=400, detail="真实姓名不能为空")
    if not company:
        raise HTTPException(status_code=400, detail="工作单位不能为空")
        
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    try:
        # 安全检查，防止把管理员账户改成别人的普通账号，或者改动管理员名字
        cursor.execute("SELECT role FROM users WHERE id = ?", (user_id,))
        user = cursor.fetchone()
        if not user:
            raise HTTPException(status_code=404, detail="该用户不存在")
        if user[0] == 'admin':
            raise HTTPException(status_code=400, detail="无权修改系统管理员账号")
            
        # 查重
        cursor.execute("SELECT id FROM users WHERE username = ? AND id != ?", (username, user_id))
        if cursor.fetchone():
            raise HTTPException(status_code=400, detail="该账号（电话）已被其他用户使用")
            
        # 更新数据
        if password and password.strip():
            cursor.execute(
                "UPDATE users SET username = ?, real_name = ?, company = ?, password = ? WHERE id = ?",
                (username, real_name, company, encrypt_pwd(password.strip()), user_id)
            )
        else:
            cursor.execute(
                "UPDATE users SET username = ?, real_name = ?, company = ? WHERE id = ?",
                (username, real_name, company, user_id)
            )
        conn.commit()
        return {"code": 200, "message": "用户信息修改成功！"}
    except Exception as e:
        if isinstance(e, HTTPException):
            raise e
        print(f"[Error] 管理员更新用户失败: {e}")
        raise HTTPException(status_code=500, detail="用户信息更新失败")
    finally:
        conn.close()

# 管理员快速修改培训人员的区域权限
@app.post("/api/admin/record/update_region")
def admin_update_region(
    record_id: int = Form(...),
    region_auth: str = Form(...),
    admin = Depends(get_admin_user)
):
    if not region_auth.strip():
        raise HTTPException(status_code=400, detail="区域权限不能为空")
        
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    # 检查记录是否存在
    cursor.execute("SELECT id FROM records WHERE id = ?", (record_id,))
    if not cursor.fetchone():
        conn.close()
        raise HTTPException(status_code=404, detail="记录不存在")
        
    cursor.execute("UPDATE records SET region_auth = ? WHERE id = ?", (region_auth.strip(), record_id))
    conn.commit()
    conn.close()
    return {"code": 200, "message": "区域权限更改成功"}

@app.post("/api/ocr_idcard")
async def api_ocr_idcard(file: UploadFile = File(...)):
    temp_ids_dir = "uploads/temp_ids"
    os.makedirs(temp_ids_dir, exist_ok=True)
    
    ext = os.path.splitext(file.filename)[1].lower()
    if not ext:
        ext = '.jpg'
    temp_path = os.path.join(temp_ids_dir, f"ocr_raw_{uuid.uuid4().hex}{ext}").replace('\\', '/')
    try:
        with open(temp_path, "wb") as f:
            shutil.copyfileobj(file.file, f)
            
        try:
            # 优先使用 Celery 异步队列 (保障并发与稳定性)
            from app.tasks import ocr_idcard_task
            task = ocr_idcard_task.delay(temp_path)
            result = task.get(timeout=8)
        except Exception as celery_err:
            print(f"[OCR-Fallback] Celery/Redis 队列异常，启用本地降级运行: {celery_err}")
            # 降级：直接在主线程中调用原有的 OCR 处理逻辑
            result = ocr_idcard_process(temp_path)
            
        if os.path.exists(temp_path):
            try: os.remove(temp_path)
            except: pass
            
        if result and "id_number" in result and "id_card" not in result:
            result["id_card"] = result["id_number"]
            
        return {"code": 200, **result}
    except Exception as e:
        if os.path.exists(temp_path):
            try: os.remove(temp_path)
            except: pass
        print(f"[Error] OCR 识别失败: {e}")
        raise HTTPException(status_code=500, detail="身份证识别失败，请重试或手动输入")

# 新接口一：提交异步 OCR 任务
@app.post("/idcard/ocr")
async def submit_ocr_task(file: UploadFile = File(...)):
    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in ['.jpg', '.jpeg', '.png', '.bmp']:
        raise HTTPException(status_code=400, detail="Unsupported image format. Allowed: jpg, jpeg, png, bmp")
        
    temp_ids_dir = "uploads/temp_ids"
    os.makedirs(temp_ids_dir, exist_ok=True)
    temp_filename = f"ocr_raw_{uuid.uuid4().hex}{ext}"
    temp_path = os.path.join(temp_ids_dir, temp_filename).replace('\\', '/')
    
    try:
        with open(temp_path, "wb") as f:
            shutil.copyfileobj(file.file, f)
            
        from app.tasks import ocr_idcard_task
        task = ocr_idcard_task.delay(temp_path)
        return {"task_id": task.id}
    except Exception as e:
        if os.path.exists(temp_path):
            try: os.remove(temp_path)
            except: pass
        raise HTTPException(status_code=500, detail=f"Failed to submit task: {str(e)}")

# 新接口二：查询异步 OCR 任务结果
@app.get("/idcard/result/{task_id}")
async def get_ocr_result(task_id: str):
    from celery.result import AsyncResult
    from app.celery_app import celery_app
    res = AsyncResult(task_id, app=celery_app)
    
    if res.status == 'SUCCESS':
        return {
            "status": "success",
            "data": res.result
        }
    elif res.status in ['PENDING', 'RECEIVED', 'STARTED', 'RETRY']:
        return {
            "status": "pending"
        }
    else:
        return {
            "status": "failed",
            "error": str(res.result) if res.result else "Unknown task execution failure"
        }

@app.get("/api/record/download_word/{record_id}")
def download_word(record_id: int, token: str = None, authorization: str = Header(None)):
    auth_token = token or authorization
    if not auth_token:
        raise HTTPException(status_code=401, detail="请重新登录")

    user_info = verify_token(auth_token)
    if not user_info:
        raise HTTPException(status_code=401, detail="请重新登录")
        
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    
    if user_info['role'] == 'admin':
        cursor.execute("""
            SELECT r.*, COALESCE(NULLIF(r.company, ''), u.company, '') AS training_company
            FROM records r 
            LEFT JOIN users u ON r.user_id = u.id 
            WHERE r.id = ?
        """, (record_id,))
    else:
        cursor.execute("""
            SELECT r.*, COALESCE(NULLIF(r.company, ''), u.company, '') AS training_company
            FROM records r 
            LEFT JOIN users u ON r.user_id = u.id 
            WHERE r.id = ? AND r.user_id = ?
        """, (record_id, user_info['id']))
        
    row = cursor.fetchone()
    conn.close()
    
    if not row:
        raise HTTPException(status_code=404, detail="登记卡未找到或您无权查看")
        
    # 无论何时，如果身份证裁剪图存在，则在下载时重新实时生成 Word 文件，以保持日期为当前下载日期，且同步所有最新文字修改
    word_path = row['word_path']
    id_card = row['id_card']
    idcard_save_dir = "uploads/idcards"
    perm_id_img_path = os.path.join(idcard_save_dir, f"{id_card}.png").replace('\\', '/')
    
    if id_card and os.path.exists(perm_id_img_path):
        try:
            record_data = {
                "姓名": row['name'],
                "性别": row['gender'],
                "年龄": row['age'],
                "联系电话": row['phone'],
                "岗位": row['job'],
                "常住地址": row['address'],
                "工作单位": row['training_company'] or '',
                "created_at": row['created_at']
            }
            temp_word_path = generate_record_card(record_data, perm_id_img_path)
            
            # 更新数据库路径
            conn_write = sqlite3.connect(DB_PATH)
            cursor_write = conn_write.cursor()
            cursor_write.execute("UPDATE records SET word_path = ? WHERE id = ?", (temp_word_path, record_id))
            conn_write.commit()
            conn_write.close()
            
            # 尝试删除旧的 Word 文件
            if word_path and word_path != temp_word_path and os.path.exists(word_path):
                try: os.remove(word_path)
                except: pass
                
            word_path = temp_word_path
        except Exception as e:
            print(f"Error regenerating word card on download: {e}")
            
    if not word_path or not os.path.exists(word_path):
        raise HTTPException(status_code=404, detail="该登记卡 Word 文件不存在或已被删除（仅保存一周）")
        
    raw_name = row['name'] or "用户"
    download_filename = f"{raw_name}登记卡.docx"
    headers = {
        "Content-Disposition": f"attachment; filename*=utf-8''{urllib.parse.quote(download_filename)}"
    }
    return FileResponse(
        word_path, 
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers
    )

# 录入培训数据
@app.post("/api/record")
async def create_record(
    name: str = Form(...),
    nation: str = Form(...),
    id_card: str = Form(...),
    phone: str = Form(...),
    address: str = Form(...),
    job: str = Form(...),
    education: str = Form(...),
    region_auth: str = Form(""),
    remark: str = Form(""),
    id_card_img_path: str = Form(None),
    photo: UploadFile = File(...),
    current_user = Depends(get_current_user)
):
    if not validate_id_card(id_card):
        raise HTTPException(status_code=400, detail="身份证号码格式不正确")

    remark = (remark or '').strip()
    if len(remark) > 1000:
        raise HTTPException(status_code=400, detail="备注不能超过1000个字符")

    # 解析身份证
    gender, age = parse_id_card(id_card)
    
    # 保存照片到 uploads
    file_ext = os.path.splitext(photo.filename)[1].lower()
    ALLOWED_EXTENSIONS = {'.jpg', '.jpeg', '.png', '.gif'}
    if not file_ext or file_ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=400, detail="不支持的图片格式，仅允许 jpg, jpeg, png, gif")
        
    # 限制大小（5MB）
    content = await photo.read(5 * 1024 * 1024 + 1)
    if len(content) > 5 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="照片文件大小不能超过 5MB")
    await photo.seek(0)
    
    filename = f"{current_user['id']}_{int(datetime.now().timestamp())}{file_ext}"
    photo_path = os.path.join(UPLOAD_DIR, filename).replace('\\', '/')
    
    with open(photo_path, "wb") as f:
        shutil.copyfileobj(photo.file, f)
        
    # 插入记录
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute('''
    INSERT INTO records (user_id, photo_path, name, nation, id_card, phone, address, job, education, region_auth, gender, age, company, remark, created_at)
    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    ''', (current_user['id'], photo_path, name, nation, id_card, phone, address, job, education, region_auth, gender, age, current_user['company'].strip(), remark, beijing_now().strftime("%Y-%m-%d %H:%M:%S")))
    record_id = cursor.lastrowid
    conn.commit()
    conn.close()
    
    # 自动生成登记卡 Word
    word_path = None
    perm_id_img_path = None
    if safe_temp_id_path(id_card_img_path):
        try:
            # 持久化身份证照片
            idcard_save_dir = "uploads/idcards"
            os.makedirs(idcard_save_dir, exist_ok=True)
            perm_id_img_path = os.path.join(idcard_save_dir, f"{id_card}.png").replace('\\', '/')
            shutil.copy2(id_card_img_path, perm_id_img_path)
            
            record_data = {
                "姓名": name,
                "性别": gender,
                "年龄": age,
                "联系电话": phone,
                "岗位": job,
                "常住地址": address,
                "工作单位": current_user['company'],
                "created_at": beijing_now().strftime("%Y-%m-%d %H:%M:%S")
            }
            word_path = generate_record_card(record_data, perm_id_img_path)
            
            # 删除临时裁剪出的身份证照片
            try: os.remove(id_card_img_path)
            except: pass
            
            # 写入数据库记录
            conn = sqlite3.connect(DB_PATH)
            cursor = conn.cursor()
            cursor.execute("UPDATE records SET word_path = ? WHERE id = ?", (word_path, record_id))
            conn.commit()
            conn.close()
        except Exception as e:
            print(f"Error generating word card: {e}")
            
    return {"code": 200, "message": "信息录入成功！"}

# 获取当前用户录入的历史记录（默认最近10天，支持姓名搜索查全局，仅录入员自己）
@app.get("/api/records")
def get_user_records(name: str = None, current_user = Depends(get_current_user)):
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    
    conditions = ["user_id = ?"]
    params = [current_user['id']]
    
    if name and name.strip():
        conditions.append("name LIKE ?")
        params.append(f"%{name.strip()}%")
    else:
        ten_days_ago = (beijing_now() - timedelta(days=9)).strftime("%Y-%m-%d 00:00:00")
        conditions.append("created_at >= ?")
        params.append(ten_days_ago)
        
    query = f'''
    SELECT * FROM records 
    WHERE {" AND ".join(conditions)}
    ORDER BY created_at DESC
    '''
    cursor.execute(query, params)
    records = [dict(row) for row in cursor.fetchall()]
    conn.close()
    return {"code": 200, "data": records}

# 修改已录入的信息（仅能修改属于自己的记录）
@app.post("/api/record/update")
async def update_record(
    record_id: int = Form(...),
    name: str = Form(...),
    nation: str = Form(...),
    id_card: str = Form(...),
    phone: str = Form(...),
    address: str = Form(...),
    job: str = Form(...),
    education: str = Form(...),
    region_auth: str = Form(""),
    remark: str = Form(""),
    id_card_img_path: str = Form(None),
    photo: UploadFile = File(None),
    current_user = Depends(get_current_user)
):
    if not validate_id_card(id_card):
        raise HTTPException(status_code=400, detail="身份证号码格式不正确")

    remark = (remark or '').strip()
    if len(remark) > 1000:
        raise HTTPException(status_code=400, detail="备注不能超过1000个字符")

    # 解析身份证
    gender, age = parse_id_card(id_card)

    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    
    # 确认记录存在且属于当前用户
    cursor.execute("SELECT * FROM records WHERE id = ? AND user_id = ?", (record_id, current_user['id']))
    record = cursor.fetchone()
    if not record:
        conn.close()
        raise HTTPException(status_code=403, detail="无权修改此记录，或记录不存在")
        
    # 判断是否创建超过一星期（7天）
    created_dt = datetime.strptime(record['created_at'], "%Y-%m-%d %H:%M:%S")
    is_over_week = (beijing_now() - created_dt) > timedelta(days=7)
    
    if is_over_week:
        # 1. 检查是否已经存在 pending 的申请
        cursor.execute("SELECT id FROM record_updates WHERE record_id = ? AND status = 'pending'", (record_id,))
        if cursor.fetchone():
            conn.close()
            raise HTTPException(status_code=400, detail="您的修改申请正在审批中，请勿重复提交")
            
        # 2. 如果用户上传了新照片，保存为临时文件以供审批
        temp_photo_path = record['photo_path']
        if photo and photo.filename:
            file_ext = os.path.splitext(photo.filename)[1].lower()
            if not file_ext:
                file_ext = '.jpg'
            filename = f"temp_update_record_{record_id}_{int(datetime.now().timestamp())}{file_ext}"
            temp_photo_path = os.path.join(UPLOAD_DIR, filename).replace('\\', '/')
            with open(temp_photo_path, "wb") as f:
                shutil.copyfileobj(photo.file, f)
        
        # 3. 身份证照临时裁剪路径暂存
        temp_id_card_img_path = None
        if safe_temp_id_path(id_card_img_path):
            temp_id_card_img_path = id_card_img_path
            
        # 4. 插入修改申请表
        cursor.execute('''
        INSERT INTO record_updates (
            record_id, user_id, name, nation, id_card, phone, address, job, education, region_auth, remark, photo_path, id_card_img_path, status, created_at
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 'pending', ?)
        ''', (
            record_id, current_user['id'], name.strip(), nation.strip(), id_card.strip(), phone.strip(),
            address.strip(), job.strip(), education.strip(), region_auth.strip(), remark, temp_photo_path, temp_id_card_img_path,
            beijing_now().strftime("%Y-%m-%d %H:%M:%S")
        ))
        conn.commit()
        conn.close()
        return {"code": 200, "message": "信息修改申请已提交，等待管理员审批", "pending_approval": True}

    photo_path = record['photo_path']
    if photo and photo.filename:
        # 保存新照片
        file_ext = os.path.splitext(photo.filename)[1]
        if not file_ext:
            file_ext = '.jpg'
        filename = f"{current_user['id']}_{int(datetime.now().timestamp())}{file_ext}"
        new_photo_path = os.path.join(UPLOAD_DIR, filename).replace('\\', '/')
        
        with open(new_photo_path, "wb") as f:
            shutil.copyfileobj(photo.file, f)
            
        # 尝试删除旧照片
        if photo_path and os.path.exists(photo_path):
            try:
                os.remove(photo_path)
            except Exception:
                pass
        photo_path = new_photo_path
        
    # 原本的 word 路径
    old_word_path = record['word_path']
    new_word_path = old_word_path
    
    idcard_save_dir = "uploads/idcards"
    os.makedirs(idcard_save_dir, exist_ok=True)
    perm_id_img_path = os.path.join(idcard_save_dir, f"{id_card}.png").replace('\\', '/')
    
    # 1. 检查是否有新上传并裁剪的身份证
    if safe_temp_id_path(id_card_img_path):
        try:
            # 覆盖原永久照片
            shutil.copy2(id_card_img_path, perm_id_img_path)
            # 删除临时照片
            try: os.remove(id_card_img_path)
            except: pass
        except Exception as e:
            print(f"Error copying new idcard image: {e}")
            
    # 2. 如果永久照片存在，重新生成 Word 登记卡
    if os.path.exists(perm_id_img_path):
        try:
            record_data = {
                "姓名": name,
                "性别": gender,
                "年龄": age,
                "联系电话": phone,
                "岗位": job,
                "常住地址": address,
                "工作单位": record['company'] or current_user['company'],
                "created_at": record['created_at']
            }
            new_word_path = generate_record_card(record_data, perm_id_img_path)
            
            # 如果生成了新的 Word 且路径发生变化，删除旧的 Word
            if old_word_path and new_word_path != old_word_path and os.path.exists(old_word_path):
                try: os.remove(old_word_path)
                except: pass
        except Exception as e:
            print(f"Error regenerating word card: {e}")
            
    cursor.execute('''
    UPDATE records 
    SET name = ?, nation = ?, id_card = ?, phone = ?, address = ?, job = ?, education = ?, region_auth = ?, remark = ?, gender = ?, age = ?, photo_path = ?, word_path = ?
    WHERE id = ? AND user_id = ?
    ''', (name, nation, id_card, phone, address, job, education, region_auth, remark, gender, age, photo_path, new_word_path, record_id, current_user['id']))
    
    conn.commit()
    conn.close()
    return {"code": 200, "message": "信息修改成功！"}

# 管理员修改培训人员记录（无权属检查）
@app.post("/api/admin/record/update")
async def admin_update_record(
    record_id: int = Form(...),
    name: str = Form(...),
    nation: str = Form(...),
    id_card: str = Form(...),
    phone: str = Form(...),
    address: str = Form(...),
    job: str = Form(...),
    education: str = Form(...),
    region_auth: str = Form(""),
    remark: str = Form(""),
    id_card_img_path: str = Form(None),
    photo: UploadFile = File(None),
    admin = Depends(get_admin_user)
):
    if not validate_id_card(id_card):
        raise HTTPException(status_code=400, detail="身份证号码格式不正确")

    remark = (remark or '').strip()
    if len(remark) > 1000:
        raise HTTPException(status_code=400, detail="备注不能超过1000个字符")

    # 解析身份证
    gender, age = parse_id_card(id_card)

    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    
    # 确认记录存在
    cursor.execute("SELECT * FROM records WHERE id = ?", (record_id,))
    record = cursor.fetchone()
    if not record:
        conn.close()
        raise HTTPException(status_code=404, detail="该人员记录不存在")
        
    photo_path = record['photo_path']
    if photo and photo.filename:
        # 保存新照片
        file_ext = os.path.splitext(photo.filename)[1]
        if not file_ext:
            file_ext = '.jpg'
        filename = f"{record['user_id']}_{int(datetime.now().timestamp())}{file_ext}"
        new_photo_path = os.path.join(UPLOAD_DIR, filename).replace('\\', '/')
        
        with open(new_photo_path, "wb") as f:
            shutil.copyfileobj(photo.file, f)
            
        # 尝试删除旧照片
        if photo_path and os.path.exists(photo_path):
            try:
                os.remove(photo_path)
            except Exception:
                pass
        photo_path = new_photo_path
        
    # 原本的 word 路径
    old_word_path = record['word_path']
    new_word_path = old_word_path
    
    idcard_save_dir = "uploads/idcards"
    os.makedirs(idcard_save_dir, exist_ok=True)
    perm_id_img_path = os.path.join(idcard_save_dir, f"{id_card}.png").replace('\\', '/')
    
    # 1. 检查是否有新上传并裁剪的身份证
    if safe_temp_id_path(id_card_img_path):
        try:
            # 覆盖原永久照片
            shutil.copy2(id_card_img_path, perm_id_img_path)
            # 删除临时照片
            try: os.remove(id_card_img_path)
            except: pass
        except Exception as e:
            print(f"Error copying new idcard image: {e}")
            
    # 2. 如果永久照片存在，重新生成 Word 登记卡
    if os.path.exists(perm_id_img_path):
        try:
            company_name = record['company'] or ''
            
            record_data = {
                "姓名": name,
                "性别": gender,
                "年龄": age,
                "联系电话": phone,
                "岗位": job,
                "常住地址": address,
                "工作单位": company_name,
                "created_at": record['created_at']
            }
            new_word_path = generate_record_card(record_data, perm_id_img_path)
            
            # 如果生成了新的 Word 且路径发生变化，删除旧 of Word
            if old_word_path and new_word_path != old_word_path and os.path.exists(old_word_path):
                try: os.remove(old_word_path)
                except: pass
        except Exception as e:
            print(f"Error regenerating word card: {e}")
            
    cursor.execute('''
    UPDATE records 
    SET name = ?, nation = ?, id_card = ?, phone = ?, address = ?, job = ?, education = ?, region_auth = ?, remark = ?, gender = ?, age = ?, photo_path = ?, word_path = ?
    WHERE id = ?
    ''', (name, nation, id_card, phone, address, job, education, region_auth, remark, gender, age, photo_path, new_word_path, record_id))
    
    conn.commit()
    conn.close()
    return {"code": 200, "message": "管理员修改记录成功！"}

# 获取所有已审核通过且未删除用户的单位列表（去重，仅管理员）
# 排序逻辑：按 exam_records 中最后答题完毕时间降序，最近有人答完题的单位排最前面
@app.get("/api/admin/companies")
def get_approved_companies(admin = Depends(get_admin_user)):
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    query = """
        WITH companies AS (
            SELECT company FROM users
            WHERE role != 'admin' AND status = 'approved' AND company IS NOT NULL AND company != ''
            UNION
            SELECT company FROM records
            WHERE company IS NOT NULL AND company != ''
        )
        SELECT c.company, MAX(e.created_at) AS last_exam_time
        FROM companies c
        LEFT JOIN exam_records e ON c.company = e.company
        GROUP BY c.company
        ORDER BY
            CASE WHEN MAX(e.created_at) IS NULL THEN 1 ELSE 0 END ASC,
            MAX(e.created_at) DESC,
            c.company ASC
    """
    cursor.execute(query)
    sorted_companies = [row[0] for row in cursor.fetchall()]
    conn.close()
    return {"code": 200, "data": sorted_companies}

# 获取所有已审核通过用户的单位列表（公开接口，供答题页面使用）
# 排序逻辑：按 exam_records 中最后答题完毕时间降序，最近有人答完题的单位排最前面
@app.get("/api/companies")
def get_companies():
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    query = """
        WITH companies AS (
            SELECT company FROM users
            WHERE role != 'admin' AND company IS NOT NULL AND company != ''
            UNION
            SELECT company FROM records
            WHERE company IS NOT NULL AND company != ''
        )
        SELECT c.company, MAX(e.created_at) AS last_exam_time
        FROM companies c
        LEFT JOIN exam_records e ON c.company = e.company
        GROUP BY c.company
        ORDER BY
            CASE WHEN MAX(e.created_at) IS NULL THEN 1 ELSE 0 END ASC,
            MAX(e.created_at) DESC,
            c.company ASC
    """
    cursor.execute(query)
    sorted_companies = [row[0] for row in cursor.fetchall()]
    conn.close()
    return {"code": 200, "data": sorted_companies}

# 查看所有已录入的信息（仅管理员，支持按日期区间筛选、工作单位筛选和门禁下载状态排序）
@app.get("/api/admin/records")
def get_all_records(start_date: str = None, end_date: str = None, company: str = None, name: str = None, page: int = 1, limit: int = 20, admin = Depends(get_admin_user)):
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    
    conditions = []
    params = []
    
    if company and company.strip():
        conditions.append("r.company = ?")
        params.append(company.strip())
        
    if name and name.strip():
        conditions.append("r.name LIKE ?")
        params.append(f"%{name.strip()}%")
    
    start = start_date.strip() if start_date and start_date.strip() else None
    end = end_date.strip() if end_date and end_date.strip() else None
    
    if start:
        conditions.append("substr(r.created_at, 1, 10) >= ?")
        params.append(start)
        
    if end:
        conditions.append("substr(r.created_at, 1, 10) <= ?")
        params.append(end)
        
    where_clause = ""
    if conditions:
        where_clause = "WHERE " + " AND ".join(conditions)
        
    total = 0
    if limit > 0:
        count_query = f'''
        SELECT COUNT(*) 
        FROM records r 
        LEFT JOIN users u ON r.user_id = u.id 
        {where_clause}
        '''
        cursor.execute(count_query, params)
        total = cursor.fetchone()[0]
        
        query = f'''
        SELECT r.*, u.real_name as recorder_name
        FROM records r 
        LEFT JOIN users u ON r.user_id = u.id 
        {where_clause}
        ORDER BY r.is_gate_downloaded ASC, r.created_at DESC
        LIMIT ? OFFSET ?
        '''
        offset = (page - 1) * limit
        cursor.execute(query, params + [limit, offset])
    else:
        query = f'''
        SELECT r.*, u.real_name as recorder_name
        FROM records r 
        LEFT JOIN users u ON r.user_id = u.id 
        {where_clause}
        ORDER BY r.is_gate_downloaded ASC, r.created_at DESC
        '''
        cursor.execute(query, params)
        
    records = [dict(row) for row in cursor.fetchall()]
    conn.close()
    return {
        "code": 200, 
        "data": records,
        "total": total if limit > 0 else len(records),
        "page": page,
        "limit": limit
    }

# 门禁导出 - 仅 CSV 导入表 (并更新已下载状态)
@app.get("/api/admin/export/gate/csv")
def export_gate_csv(ids: str = None, admin = Depends(get_admin_user)):
    id_list = parse_id_list(ids)
    with get_db() as conn:
        cursor = conn.cursor()
        records, placeholders = fetch_records_by_ids(cursor, id_list, gate_only=False)
        if not records:
            raise HTTPException(status_code=404, detail="未找到对应的记录")
        csv_data = build_gate_csv(records)
        # 将门禁下载状态更新为已下载(1)
        cursor.execute(f"UPDATE records SET is_gate_downloaded = 1 WHERE id IN ({placeholders})", id_list)
        conn.commit()
    return _csv_download_response(csv_data, "培训人员导入表.csv")

# 门禁导出 - 仅照片压缩包
@app.get("/api/admin/export/gate/photos")
def export_gate_photos(ids: str = None, admin = Depends(get_admin_user)):
    id_list = parse_id_list(ids)
    with get_db() as conn:
        cursor = conn.cursor()
        records, _ = fetch_records_by_ids(cursor, id_list, gate_only=False)
    if not records:
        raise HTTPException(status_code=404, detail="未找到对应的记录")
    return _zip_download_response(pack_photos_zip(records), "培训人员照片.zip")

# 门禁下载旧版兼容接口 (CSV 导入表和照片打包，支持按 ids 筛选，并更新已下载状态)
@app.get("/api/admin/export/gate")
def export_gate_old_compatible(ids: str = None, admin = Depends(get_admin_user)):
    id_list = parse_id_list(ids)
    with get_db() as conn:
        cursor = conn.cursor()
        records, placeholders = fetch_records_by_ids(cursor, id_list, gate_only=False)
        if not records:
            raise HTTPException(status_code=404, detail="未找到对应的记录")
        csv_data = build_gate_csv(records)
        # CSV 导入表 + 照片合并打包
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            zip_file.writestr("培训人员导入表.csv", csv_data)
            added_filenames = set()
            for r in records:
                photo_path = r['photo_path']
                if not photo_path or not os.path.exists(photo_path):
                    continue
                file_ext = os.path.splitext(photo_path)[1] or '.jpg'
                base_filename = f"{safe_filename_part(r['name'])}_{safe_filename_part(r['id_card'])}"
                filename = f"{base_filename}{file_ext}"
                counter = 1
                while filename in added_filenames:
                    filename = f"{base_filename}_{counter}{file_ext}"
                    counter += 1
                added_filenames.add(filename)
                zip_file.write(photo_path, arcname=filename)
        cursor.execute(f"UPDATE records SET is_gate_downloaded = 1 WHERE id IN ({placeholders})", id_list)
        conn.commit()
    return _zip_download_response(zip_buffer.getvalue(), "门禁系统导入包.zip")

# 导出并下载 Excel (按已有模板的格式，支持按 ids/日期区间 筛选)
@app.get("/api/admin/export/excel")
def export_excel(ids: str = None, start_date: str = None, end_date: str = None, company: str = None, name: str = None, admin = Depends(get_admin_user)):
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    
    conditions = []
    params = []
    
    if company and company.strip():
        conditions.append("r.company = ?")
        params.append(company.strip())
        
    if name and name.strip():
        conditions.append("r.name LIKE ?")
        params.append(f"%{name.strip()}%")
    
    if ids:
        id_list = [int(x) for x in ids.split(',') if x.strip()]
        placeholders = ','.join(['?'] * len(id_list))
        conditions.append(f"r.id IN ({placeholders})")
        params.extend(id_list)
        
    start = start_date.strip() if start_date and start_date.strip() else None
    end = end_date.strip() if end_date and end_date.strip() else None
    
    if start and not end:
        end = start
    elif end and not start:
        start = end
        
    if start:
        conditions.append("substr(r.created_at, 1, 10) >= ?")
        params.append(start)
        
    if end:
        conditions.append("substr(r.created_at, 1, 10) <= ?")
        params.append(end)
        
    where_clause = ""
    if conditions:
        where_clause = "WHERE " + " AND ".join(conditions)
        
    query = f'''
    SELECT r.*
    FROM records r 
    LEFT JOIN users u ON r.user_id = u.id 
    {where_clause}
    ORDER BY r.is_gate_downloaded ASC, r.created_at DESC
    '''
    
    cursor.execute(query, params)
    records = cursor.fetchall()
    
    template_path = '06.03.xlsx'
    if not os.path.exists(template_path):
        # 如果模板不见了，则重新生成一个空的带表头的 Workbook
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Sheet1"
        ws.append([
            "姓名", "性别", "民族", "年龄", "身份证号码", "联系电话", "现常住地址", "公司名称", 
            "岗位/工种", "学历", "培训日期", "考试成绩", "有效期限", "区域权限", "人员在各单位间流动情况", 
            "最近一次培训日期", "特殊工种证有效期", "备注"
        ])
    else:
        wb = openpyxl.load_workbook(template_path)
        ws = wb.active
        
    # 清空除第一行表头外的所有行（以防模板里有之前的数据）
    if ws.max_row > 1:
        ws.delete_rows(2, ws.max_row)
        
    for r in records:
        # 查询最近一次答题记录
        r_company = r['company'] if r['company'] else ""
        cursor.execute('''
            SELECT created_at, score FROM exam_records 
            WHERE name = ? AND company = ? 
            ORDER BY created_at DESC LIMIT 1
        ''', (r['name'], r_company))
        exam_row = cursor.fetchone()
        
        train_date = ""
        exam_score = ""
        if exam_row:
            train_date = exam_row[0][:10] if exam_row[0] else ""  # 取 YYYY-MM-DD
            exam_score = exam_row[1]
            
        row_data = [
            r['name'],
            r['gender'],
            r['nation'],
            r['age'],
            r['id_card'],
            r['phone'],
            r['address'],
            r['company'],
            r['job'],
            r['education'],
            train_date, # 培训日期
            exam_score, # 考试成绩
            "", # 有效期限
            r['region_auth'], # 区域权限
            "", # 流动情况
            "", # 最近一次培训日期
            "", # 特殊工种证有效期
            r['remark'] or ""  # 备注
        ]
        ws.append(row_data)
        
    conn.close()

    # L5: 用内存缓冲返回，避免写死文件名在并发导出时互相覆盖
    out_buffer = io.BytesIO()
    wb.save(out_buffer)
    out_buffer.seek(0)
    return Response(
        content=out_buffer.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename*=utf-8''%E5%9F%B9%E8%AE%AD%E4%BA%BA%E5%91%98%E4%BF%A1%E6%81%AF%E8%A1%A8.xlsx"}
    )

# 导出并下载 CSV (按已有模板的格式)
@app.get("/api/admin/export/csv")
def export_csv(admin = Depends(get_admin_user)):
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute('''
        SELECT r.*
        FROM records r
        LEFT JOIN users u ON r.user_id = u.id
        ORDER BY r.created_at DESC
        ''')
        records = cursor.fetchall()
    return _csv_download_response(build_gate_csv(records), "培训人员导入表.csv")

# 导出并下载照片压缩包（仅管理员）
@app.get("/api/admin/export/photos")
def export_photos(admin = Depends(get_admin_user)):
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT photo_path, name, id_card FROM records")
        records = cursor.fetchall()
    return _zip_download_response(pack_photos_zip(records), "培训人员照片.zip")

# 获取同单位的所有人员记录（供客户端检索并用于门禁恢复）
@app.get("/api/user/company_records")
def get_company_records(name: str = None, current_user = Depends(get_current_user)):
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    
    query = '''
    SELECT r.*, u.real_name as recorder_name
    FROM records r
    LEFT JOIN users u ON r.user_id = u.id
    WHERE r.company = ? AND (r.gate_restore_status IS NULL OR r.gate_restore_status != 'pending')
    '''
    params = [current_user['company']]
    
    if name and name.strip():
        query += " AND r.name LIKE ?"
        params.append(f"%{name.strip()}%")
        
    query += " ORDER BY r.created_at DESC"
    
    cursor.execute(query, params)
    records = [dict(row) for row in cursor.fetchall()]
    conn.close()
    return {"code": 200, "data": records}

# 提交门禁恢复申请（回复门禁）
@app.post("/api/user/record/restore_gate")
def restore_gate(ids: str = Form(...), current_user = Depends(get_current_user)):
    if not ids:
        raise HTTPException(status_code=400, detail="请选择需要恢复门禁的人员")
        
    id_list = [int(x) for x in ids.split(',') if x.strip()]
    if not id_list:
        raise HTTPException(status_code=400, detail="没有勾选任何有效的人员记录")
        
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    # 验证这些 records 都属于该用户所在的公司，防越权
    placeholders = ','.join(['?'] * len(id_list))
    check_query = f'''
    SELECT r.id FROM records r
    LEFT JOIN users u ON r.user_id = u.id
    WHERE r.id IN ({placeholders}) AND r.company = ?
    '''
    cursor.execute(check_query, id_list + [current_user['company']])
    allowed_ids = [row[0] for row in cursor.fetchall()]
    
    if len(allowed_ids) != len(id_list):
        conn.close()
        raise HTTPException(status_code=403, detail="部分记录不存在或无权操作")
        
    # 更新 records 的门禁恢复状态
    update_query = f'''
    UPDATE records
    SET gate_restore_status = 'pending', is_restore_downloaded = 0
    WHERE id IN ({placeholders})
    '''
    cursor.execute(update_query, id_list)
    conn.commit()
    conn.close()
    return {"code": 200, "message": "门禁恢复申请提交成功"}

# 查看所有提交了门禁恢复申请的人员记录（仅管理员）
@app.get("/api/admin/restore_records")
def get_restore_records(admin = Depends(get_admin_user)):
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    
    cursor.execute('''
    SELECT r.*, u.real_name as recorder_name
    FROM records r
    LEFT JOIN users u ON r.user_id = u.id
    WHERE r.gate_restore_status = 'pending'
    ORDER BY r.is_restore_downloaded ASC, r.created_at DESC
    ''')
    records = [dict(row) for row in cursor.fetchall()]
    conn.close()
    return {"code": 200, "data": records}

# 删除门禁恢复申请记录（仅管理员）
@app.post("/api/admin/delete_restore_gate")
def delete_restore_gate(ids: str = Form(...), admin = Depends(get_admin_user)):
    if not ids:
        raise HTTPException(status_code=400, detail="请选择需要删除的记录")
        
    id_list = [int(x) for x in ids.split(',') if x.strip()]
    if not id_list:
        raise HTTPException(status_code=400, detail="无有效记录ID")
        
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    placeholders = ','.join(['?'] * len(id_list))
    update_query = f'''
    UPDATE records
    SET gate_restore_status = NULL, is_restore_downloaded = 0
    WHERE id IN ({placeholders})
    '''
    cursor.execute(update_query, id_list)
    conn.commit()
    conn.close()
    return {"code": 200, "message": "门禁恢复记录删除成功"}


# 门禁恢复导出 - 仅 CSV 导入表 (并更新 is_restore_downloaded)
@app.get("/api/admin/export/restore/csv")
def export_restore_csv(ids: str = None, admin = Depends(get_admin_user)):
    id_list = parse_id_list(ids)
    with get_db() as conn:
        cursor = conn.cursor()
        records, placeholders = fetch_records_by_ids(cursor, id_list, gate_only=True)
        if not records:
            raise HTTPException(status_code=404, detail="未找到对应的恢复申请记录")
        csv_data = build_gate_csv(records)
        # 将门禁恢复下载状态更新为已下载(1)
        cursor.execute(f"UPDATE records SET is_restore_downloaded = 1 WHERE id IN ({placeholders})", id_list)
        conn.commit()
    return _csv_download_response(csv_data, "恢复人员导入表.csv")

# 门禁恢复导出 - 仅照片压缩包
@app.get("/api/admin/export/restore/photos")
def export_restore_photos(ids: str = None, admin = Depends(get_admin_user)):
    id_list = parse_id_list(ids)
    with get_db() as conn:
        cursor = conn.cursor()
        records, _ = fetch_records_by_ids(cursor, id_list, gate_only=True)
    if not records:
        raise HTTPException(status_code=404, detail="未找到对应的记录")
    return _zip_download_response(pack_photos_zip(records), "恢复人员照片.zip")

# 题库上传更新
@app.post("/api/admin/upload_exam_bank")
async def upload_exam_bank(
    exam_type: str = Form(...),
    file: UploadFile = File(...),
    admin = Depends(get_admin_user)
):
    """管理员上传 xlsx 题库文件，替换 shiti/ 目录中的对应文件"""
    exam_file_map = get_exam_file_map()
    target_filename = exam_file_map.get(exam_type)
    if not target_filename:
        raise HTTPException(status_code=400, detail=f"未知的考试类型: {exam_type}")

    if not file.filename.endswith('.xlsx'):
        raise HTTPException(status_code=400, detail="仅支持 .xlsx 格式文件")

    content = await file.read()

    # 验证文件有效性
    try:
        wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True)
        ws = wb.active
        count = 0
        for row in ws.iter_rows(min_row=2, values_only=True):
            if len(row) >= 4 and row[1] and row[3]:
                count += 1
        wb.close()
        if count == 0:
            raise HTTPException(status_code=400, detail="题库文件中未找到有效题目（需B列题目+D列答案）")
    except HTTPException:
        raise
    except Exception as e:
        print(f"[Error] 题库文件解析失败: {e}")
        raise HTTPException(status_code=400, detail="题库文件格式错误，无法解析")
    
    # 保存文件
    os.makedirs('shiti', exist_ok=True)
    target_path = os.path.join('shiti', target_filename)
    with open(target_path, 'wb') as f:
        f.write(content)
    
    # 清除缓存
    if exam_type in EXAM_QUESTIONS_CACHE:
        del EXAM_QUESTIONS_CACHE[exam_type]
    
    return {"code": 200, "message": f"{exam_type}题库更新成功", "question_count": count}

# ---------------- 考试相关接口 ----------------

# 考试身份验证（无需登录）
@app.post("/api/exam/verify")
def exam_verify(data: dict):
    name = (data.get('name') or '').strip()
    id_last6 = (data.get('id_last6') or '').strip()
    
    if not name or not id_last6 or len(id_last6) != 6:
        raise HTTPException(status_code=400, detail="请输入姓名和身份证号后六位")
    
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    
    # Match by both name and ID suffix, then take the most recently uploaded
    # training entry.  Older training entries stay available in search results.
    latest_training = get_latest_training_record(cursor, name, id_last6)
    if latest_training:
        conn.close()
        return {
            "code": 200,
            "matched": True,
            "company": latest_training['training_company'] or ''
        }

    cursor.execute("SELECT 1 FROM records WHERE SUBSTR(id_card, -6) = ? LIMIT 1", (id_last6,))
    if cursor.fetchone():
        # 身份证后六位匹配但姓名不匹配
        conn.close()
        return {"code": 200, "matched": False, "error": "name_mismatch"}
    
    # 身份证后六位不匹配，检查姓名是否存在
    cursor.execute("SELECT id FROM records WHERE name = ?", (name,))
    name_matches = cursor.fetchall()
    
    if name_matches:
        # 姓名匹配但身份证后六位不匹配
        conn.close()
        return {"code": 200, "matched": False, "error": "id_mismatch"}
    
    # 都不匹配，检查审批表
    cursor.execute("""
        SELECT * FROM exam_approvals
        WHERE name = ? AND id_last6 = ?
        ORDER BY created_at DESC LIMIT 1
    """, (name, id_last6))
    approval = cursor.fetchone()
    
    if approval:
        if approval['status'] == 'approved':
            conn.close()
            return {"code": 200, "matched": True, "company": approval['company'] or '', "approved": True}
        elif approval['status'] == 'pending':
            conn.close()
            return {"code": 200, "matched": False, "error": "pending_approval"}
        elif approval['status'] == 'rejected':
            # 被拒绝后可以重新提交
            cursor.execute("""
                INSERT INTO exam_approvals (name, id_last6, status, created_at)
                VALUES (?, ?, 'pending', ?)
            """, (name, id_last6, beijing_now().strftime("%Y-%m-%d %H:%M:%S")))
            conn.commit()
            conn.close()
            return {"code": 200, "matched": False, "error": "not_found"}
    
    # 完全没有记录，创建审批申请
    cursor.execute("""
        INSERT INTO exam_approvals (name, id_last6, status, created_at)
        VALUES (?, ?, 'pending', ?)
    """, (name, id_last6, beijing_now().strftime("%Y-%m-%d %H:%M:%S")))
    conn.commit()
    conn.close()
    return {"code": 200, "matched": False, "error": "not_found"}

# 获取答题审批列表（管理员）
@app.get("/api/admin/exam-approvals")
def get_exam_approvals(admin = Depends(get_admin_user)):
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM exam_approvals ORDER BY CASE WHEN status='pending' THEN 0 ELSE 1 END, created_at DESC")
    approvals = [dict(row) for row in cursor.fetchall()]
    conn.close()
    return {"code": 200, "data": approvals}

# 审批答题申请（管理员）
@app.put("/api/admin/exam-approvals/{approval_id}/approve")
def approve_exam_approval(approval_id: int, data: dict = None, admin = Depends(get_admin_user)):
    company = ''
    if data:
        company = (data.get('company') or '').strip()
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("UPDATE exam_approvals SET status = 'approved', company = ? WHERE id = ?", (company, approval_id))
    if cursor.rowcount == 0:
        conn.close()
        raise HTTPException(status_code=404, detail="审批记录不存在")
    conn.commit()
    conn.close()
    return {"code": 200, "message": "已批准该答题申请"}

# 拒绝答题申请（管理员）
@app.put("/api/admin/exam-approvals/{approval_id}/reject")
def reject_exam_approval(approval_id: int, admin = Depends(get_admin_user)):
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("UPDATE exam_approvals SET status = 'rejected' WHERE id = ?", (approval_id,))
    if cursor.rowcount == 0:
        conn.close()
        raise HTTPException(status_code=404, detail="审批记录不存在")
    conn.commit()
    conn.close()
    return {"code": 200, "message": "已拒绝该答题申请"}

# ================== 新增：人员信息修改审批 & 消息提醒 API ==================

# 1. 获取修改申请列表（管理员）
@app.get("/api/admin/record-updates")
def get_record_updates(admin = Depends(get_admin_user)):
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    cursor.execute('''
        SELECT 
            ru.id as update_id, ru.record_id, ru.user_id, ru.status, ru.created_at as apply_time,
            ru.name as new_name, ru.nation as new_nation, ru.id_card as new_id_card, ru.phone as new_phone,
            ru.address as new_address, ru.job as new_job, ru.education as new_education, ru.region_auth as new_region_auth, ru.remark as new_remark,
            ru.photo_path as new_photo_path, ru.id_card_img_path as new_id_card_img_path,
            r.name as old_name, r.nation as old_nation, r.id_card as old_id_card, r.phone as old_phone,
            r.address as old_address, r.job as old_job, r.education as old_education, r.region_auth as old_region_auth, r.remark as old_remark,
            r.photo_path as old_photo_path, r.word_path as old_word_path,
            COALESCE(NULLIF(r.company, ''), u.company, '') as old_company
        FROM record_updates ru
        LEFT JOIN records r ON ru.record_id = r.id
        LEFT JOIN users u ON ru.user_id = u.id
        ORDER BY CASE WHEN ru.status='pending' THEN 0 ELSE 1 END, ru.created_at DESC
    ''')
    updates = [dict(row) for row in cursor.fetchall()]
    conn.close()
    return {"code": 200, "data": updates}

# 2. 批准人员信息修改申请（管理员）
@app.post("/api/admin/record-updates/{update_id}/approve")
def approve_record_update(update_id: int, admin = Depends(get_admin_user)):
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    
    cursor.execute("SELECT * FROM record_updates WHERE id = ?", (update_id,))
    update_req = cursor.fetchone()
    if not update_req:
        conn.close()
        raise HTTPException(status_code=404, detail="申请记录不存在")
        
    if update_req['status'] != 'pending':
        conn.close()
        raise HTTPException(status_code=400, detail="该申请已被处理")
        
    record_id = update_req['record_id']
    user_id = update_req['user_id']
    
    cursor.execute("SELECT * FROM records WHERE id = ?", (record_id,))
    record = cursor.fetchone()
    
    if not record:
        conn.close()
        raise HTTPException(status_code=404, detail="对应的培训人员记录不存在")

    company = record['company'] or ''
        
    # 处理头像照片的转正和旧照物理删除
    final_photo_path = record['photo_path']
    new_photo_req = update_req['photo_path']
    if new_photo_req and new_photo_req != record['photo_path']:
        if os.path.exists(new_photo_req):
            formal_filename = os.path.basename(new_photo_req).replace("temp_update_record_", "")
            formal_photo_path = os.path.join(UPLOAD_DIR, formal_filename).replace('\\', '/')
            try:
                os.rename(new_photo_req, formal_photo_path)
                final_photo_path = formal_photo_path
            except Exception:
                final_photo_path = new_photo_req
                
            if record['photo_path'] and os.path.exists(record['photo_path']):
                try: os.remove(record['photo_path'])
                except: pass
                
    # 处理新身份证裁剪图片的转正
    new_id_card = update_req['id_card']
    idcard_save_dir = "uploads/idcards"
    os.makedirs(idcard_save_dir, exist_ok=True)
    perm_id_img_path = os.path.join(idcard_save_dir, f"{new_id_card}.png").replace('\\', '/')
    
    temp_id_card_img_path = update_req['id_card_img_path']
    if temp_id_card_img_path and os.path.exists(temp_id_card_img_path):
        try:
            shutil.copy2(temp_id_card_img_path, perm_id_img_path)
            try: os.remove(temp_id_card_img_path)
            except: pass
        except Exception as e:
            print(f"Error copying idcard image during approval: {e}")
            
    # 解析新身份证信息并重新生成 Word 登记卡
    gender, age = parse_id_card(new_id_card)
    new_word_path = record['word_path']
    if os.path.exists(perm_id_img_path):
        try:
            record_data = {
                "姓名": update_req['name'],
                "性别": gender,
                "年龄": age,
                "联系电话": update_req['phone'],
                "岗位": update_req['job'],
                "常住地址": update_req['address'],
                "工作单位": company,
                "created_at": record['created_at']
            }
            new_word_path = generate_record_card(record_data, perm_id_img_path)
            
            if record['word_path'] and new_word_path != record['word_path'] and os.path.exists(record['word_path']):
                try: os.remove(record['word_path'])
                except: pass
        except Exception as e:
            print(f"Error regenerating word card during approval: {e}")
            
    # 更新 records 记录
    cursor.execute('''
        UPDATE records 
        SET name = ?, nation = ?, id_card = ?, phone = ?, address = ?, job = ?, education = ?, region_auth = ?, remark = ?, gender = ?, age = ?, photo_path = ?, word_path = ?
        WHERE id = ?
    ''', (
        update_req['name'], update_req['nation'], new_id_card, update_req['phone'], update_req['address'],
        update_req['job'], update_req['education'], update_req['region_auth'], update_req['remark'], gender, age, final_photo_path, new_word_path,
        record_id
    ))
    
    # 更新申请状态
    cursor.execute("UPDATE record_updates SET status = 'approved' WHERE id = ?", (update_id,))
    
    conn.commit()
    conn.close()
    return {"code": 200, "message": "审批已通过，人员信息已更新"}

# 3. 拒绝人员信息修改申请（管理员）
@app.post("/api/admin/record-updates/{update_id}/reject")
def reject_record_update(update_id: int, admin = Depends(get_admin_user)):
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    
    cursor.execute("SELECT * FROM record_updates WHERE id = ?", (update_id,))
    update_req = cursor.fetchone()
    if not update_req:
        conn.close()
        raise HTTPException(status_code=404, detail="申请记录不存在")
        
    if update_req['status'] != 'pending':
        conn.close()
        raise HTTPException(status_code=400, detail="该申请已被处理")
        
    record_id = update_req['record_id']
    cursor.execute("SELECT photo_path FROM records WHERE id = ?", (record_id,))
    record = cursor.fetchone()
    
    # 物理删除临时生成的照片
    temp_photo_path = update_req['photo_path']
    if temp_photo_path and record and temp_photo_path != record['photo_path']:
        if os.path.exists(temp_photo_path):
            try: os.remove(temp_photo_path)
            except: pass
            
    # 物理删除临时身份证照片
    temp_id_card_img_path = update_req['id_card_img_path']
    if temp_id_card_img_path and os.path.exists(temp_id_card_img_path):
        try: os.remove(temp_id_card_img_path)
        except: pass
        
    cursor.execute("UPDATE record_updates SET status = 'rejected' WHERE id = ?", (update_id,))
    conn.commit()
    conn.close()
    return {"code": 200, "message": "审批已拒绝，原人员信息保持不变"}

# 4. 获取管理员消息提醒数（答题没有数据的提醒 + 超过7天修改审批）
@app.get("/api/admin/badge-count")
def get_badge_count(admin = Depends(get_admin_user)):
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    cursor.execute("SELECT COUNT(*) FROM exam_approvals WHERE status = 'pending'")
    exam_pending = cursor.fetchone()[0]
    
    cursor.execute("SELECT COUNT(*) FROM record_updates WHERE status = 'pending'")
    record_pending = cursor.fetchone()[0]
    
    try:
        cursor.execute("SELECT COUNT(*) FROM users WHERE role != 'admin' AND status = 'pending' AND is_deleted = 0")
        user_pending = cursor.fetchone()[0]
    except sqlite3.OperationalError:
        try:
            cursor.execute("SELECT COUNT(*) FROM users WHERE role != 'admin' AND status = 'pending'")
            user_pending = cursor.fetchone()[0]
        except sqlite3.OperationalError:
            user_pending = 0
    
    conn.close()
    return {
        "code": 200,
        "data": {
            "exam_pending": exam_pending,
            "record_pending": record_pending,
            "user_pending": user_pending,
            "total": exam_pending + record_pending + user_pending
        }
    }

# =========================================================================

# 获取试题（无需登录）
@app.get("/api/get_questions")
def get_questions(file: str):
    if not is_exam_open():
        raise HTTPException(status_code=403, detail="当前非考试开放时间，禁止获取试卷")

    # S6: 白名单校验——只允许 configs 中已配置的题库文件名，杜绝路径遍历
    allowed_files = set(get_exam_file_map().values())
    if file not in allowed_files:
        raise HTTPException(status_code=400, detail="非法或未知的试题文件")

    # 构建完整路径
    shiti_dir = 'shiti'
    file_path = os.path.join(shiti_dir, file)
    
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="试题文件不存在")
    
    try:
        wb = openpyxl.load_workbook(file_path)
        ws = wb.active
        
        questions = []
        # 表头：题目类型、题目、解析、正确答案、答案A...
        # 跳过第一行表头，从第二行开始读取
        for row in ws.iter_rows(min_row=2, values_only=True):
            if len(row) >= 4 and row[1] and row[3]:
                question_text = str(row[1]).strip()
                
                # 仅返回题目，屏蔽正确答案以防前端网络拦截作弊
                question = {
                    "question": question_text
                }
                questions.append(question)
        
        return {"code": 200, "data": questions}
    except Exception as e:
        print(f"读取试题失败: {e}")
        raise HTTPException(status_code=500, detail="读取试题失败")

# 保存答题记录（防作弊后端判分）
@app.post("/api/save_exam_record")
def save_exam_record(data: dict):
    if not is_exam_open():
        raise HTTPException(status_code=403, detail="当前非考试开放时间，禁止提交答卷")
    conn = None
    try:
        name = (data.get('name') or '').strip()
        company = (data.get('company') or '').strip()
        id_last6 = (data.get('id_last6') or '').strip()
        exam_type = data.get('exam_type')
        duration = data.get('duration')
        user_answers = data.get('answers', [])  # 格式: [{"question": "xxx", "user_answer": "对/错"}]

        if not name or not exam_type or len(id_last6) != 6:
            raise HTTPException(status_code=400, detail="缺少必要信息")

        # 1. 在后端重新加载正确答案，确保防作弊安全性（使用缓存优化）
        shiti_answers = get_exam_questions_answers(exam_type)

        # 2. 对比计分
        correct_count = 0
        details = []
        for ua in user_answers:
            q_text = ua.get('question', '').strip()
            user_ans = ua.get('user_answer', '').strip()
            correct_ans = shiti_answers.get(q_text, '')

            is_correct_bool = (user_ans == correct_ans) if correct_ans else False
            if is_correct_bool:
                correct_count += 1

            details.append({
                "question": q_text,
                "user_answer": user_ans,
                "correct_answer": correct_ans,
                "is_correct": 1 if is_correct_bool else 0
            })

        total_questions = len(user_answers)
        score = round((correct_count / total_questions) * 100) if total_questions > 0 else 0
        answered_count = len([a for a in user_answers if a.get('user_answer')])

        conn = sqlite3.connect(DB_PATH)
        conn.row_factory = sqlite3.Row
        cursor = conn.cursor()

        # Never trust the company sent by the browser for an existing trainee.
        # The newest training record is the source of truth for the exam unit.
        latest_training = get_latest_training_record(cursor, name, id_last6)
        if latest_training:
            company = latest_training['training_company'] or ''
        else:
            cursor.execute('''
                SELECT 1 FROM exam_approvals
                WHERE name = ? AND id_last6 = ? AND status = 'approved'
                ORDER BY created_at DESC LIMIT 1
            ''', (name, id_last6))
            if not cursor.fetchone():
                raise HTTPException(status_code=403, detail="请先完成身份验证")

        if not company:
            raise HTTPException(status_code=400, detail="缺少工作单位")

        # 频率限制：同一姓名+单位+科目 5 分钟内只能提交一次，防恶意刷分，不影响正常重考
        cursor.execute('''
            SELECT created_at FROM exam_records
            WHERE name = ? AND company = ? AND exam_type = ?
            ORDER BY created_at DESC LIMIT 1
        ''', (name, company, exam_type))
        recent = cursor.fetchone()
        if recent and recent[0]:
            try:
                last_time = datetime.strptime(str(recent[0])[:19], "%Y-%m-%d %H:%M:%S")
                if (beijing_now() - last_time).total_seconds() < 300:
                    raise HTTPException(status_code=429, detail="提交过于频繁，请 5 分钟后再试")
            except HTTPException:
                raise
            except Exception:
                pass  # 时间解析失败则放行，不阻塞正常流程

        # 保留全部历史记录（不覆盖），重考会产生多条；查询与导出时取最新一条即可
        # （export_excel 已 ORDER BY created_at DESC LIMIT 1）。

        # 插入答题记录
        cursor.execute('''
            INSERT INTO exam_records (name, company, exam_type, score, answered_count, correct_count, duration, created_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
        ''', (name, company, exam_type, score, answered_count, correct_count, duration, beijing_now().strftime("%Y-%m-%d %H:%M:%S")))

        record_id = cursor.lastrowid

        # 插入答题详情
        for d in details:
            cursor.execute('''
                INSERT INTO exam_details (exam_record_id, question, user_answer, correct_answer, is_correct)
                VALUES (?, ?, ?, ?, ?)
            ''', (record_id, d["question"], d["user_answer"], d["correct_answer"], d["is_correct"]))

        conn.commit()
        conn.close()

        return {
            "code": 200,
            "message": "答题记录保存成功",
            "data": {
                "score": score,
                "correct_count": correct_count,
                "answered_count": answered_count,
                "total_count": total_questions
            }
        }
    except HTTPException:
        if conn:
            try:
                conn.rollback()
                conn.close()
            except Exception:
                pass
        raise  # L2: 保留 400/403/429 等具体错误，不被通用异常吞掉
    except Exception as e:
        print(f"保存答题记录失败: {e}")
        if conn:  # L7: 异常时回滚，避免产生有主记录无详情的脏数据
            try:
                conn.rollback()
                conn.close()
            except Exception:
                pass
        raise HTTPException(status_code=500, detail="保存答题记录与后端判分失败")

# 获取答题记录列表（管理员权限）
@app.get("/api/admin/exam_records")
def get_exam_records(company: str = '', exam_type: str = '', name: str = '', 
                     page: int = 1, limit: int = 20,
                     current_user = Depends(get_admin_user)):
    try:
        conn = sqlite3.connect(DB_PATH)
        conn.row_factory = sqlite3.Row
        cursor = conn.cursor()
        
        base_query = "FROM exam_records WHERE 1=1"
        params = []
        
        if company:
            base_query += " AND company LIKE ?"
            params.append(f"%{company}%")
        if exam_type:
            base_query += " AND exam_type = ?"
            params.append(exam_type)
        if name:
            base_query += " AND name LIKE ?"
            params.append(f"%{name}%")
            
        # 查出全部符合条件的记录，以便在内存中去重合并多次答题
        query = f"SELECT * {base_query} ORDER BY created_at DESC"
        cursor.execute(query, params)
        records = cursor.fetchall()
        
        grouped = {}
        for r in records:
            key = (r['name'], r['company'], r['exam_type'])
            if key not in grouped:
                grouped[key] = {
                    "id": r['id'],
                    "name": r['name'],
                    "company": r['company'],
                    "exam_type": r['exam_type'],
                    "score": r['score'],
                    "answered_count": r['answered_count'],
                    "correct_count": r['correct_count'],
                    "duration": r['duration'],
                    "created_at": r['created_at'],
                    "history": []
                }
            grouped[key]["history"].append({
                "id": r['id'],
                "score": r['score'],
                "answered_count": r['answered_count'],
                "correct_count": r['correct_count'],
                "duration": r['duration'],
                "created_at": r['created_at']
            })
            
        grouped_list = list(grouped.values())
        total = len(grouped_list)
        
        # 进行分页
        if limit > 0:
            offset = (page - 1) * limit
            paginated_records = grouped_list[offset : offset + limit]
        else:
            paginated_records = grouped_list
            
        conn.close()
        
        return {
            "code": 200, 
            "data": paginated_records,
            "total": total,
            "page": page,
            "limit": limit
        }
    except Exception as e:
        print(f"获取答题记录失败: {e}")
        raise HTTPException(status_code=500, detail="获取答题记录失败")

# 获取答题详情（管理员权限）
@app.get("/api/admin/exam_record_detail/{record_id}")
def get_exam_record_detail(record_id: int, current_user = Depends(get_admin_user)):
    try:
        conn = sqlite3.connect(DB_PATH)
        conn.row_factory = sqlite3.Row
        cursor = conn.cursor()
        
        # 获取主记录
        cursor.execute("SELECT * FROM exam_records WHERE id = ?", (record_id,))
        record = cursor.fetchone()
        
        if not record:
            raise HTTPException(status_code=404, detail="答题记录不存在")
        
        # 获取详情
        cursor.execute("SELECT * FROM exam_details WHERE exam_record_id = ? ORDER BY id", (record_id,))
        details = cursor.fetchall()
        
        detail_list = []
        for d in details:
            detail_list.append({
                "question": d['question'],
                "user_answer": d['user_answer'],
                "correct_answer": d['correct_answer'],
                "is_correct": d['is_correct']
            })
        
        conn.close()
        
        return {"code": 200, "data": {
            "id": record['id'],
            "name": record['name'],
            "company": record['company'],
            "exam_type": record['exam_type'],
            "score": record['score'],
            "answered_count": record['answered_count'],
            "correct_count": record['correct_count'],
            "duration": record['duration'],
            "created_at": record['created_at'],
            "details": detail_list
        }}
    except HTTPException:
        raise
    except Exception as e:
        print(f"获取答题详情失败: {e}")
        raise HTTPException(status_code=500, detail="获取答题详情失败")

# 导出答题过程 Word（仅管理员）
@app.get("/api/admin/exam_records/{record_id}/download")
def download_exam_record(record_id: int, admin = Depends(get_admin_user)):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io
    import urllib.parse
    from fastapi.responses import StreamingResponse
    
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    
    # 查 exam_records
    cursor.execute("SELECT * FROM exam_records WHERE id = ?", (record_id,))
    record = cursor.fetchone()
    if not record:
        conn.close()
        raise HTTPException(status_code=404, detail="未找到对应的答题记录")
        
    # 查 exam_details
    cursor.execute("SELECT * FROM exam_details WHERE exam_record_id = ? ORDER BY id", (record_id,))
    details = cursor.fetchall()
    conn.close()
    
    doc = Document()
    
    # 设置大标题
    title = doc.add_paragraph()
    title_run = title.add_run("在线答题过程及详情报告")
    title_run.font.size = Pt(16)
    title_run.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 基本信息表格
    table = doc.add_table(rows=3, cols=4)
    table.style = 'Table Grid'
    
    # 表格内容填充
    def set_cell_text(row_idx, col_idx, label, val):
        row = table.rows[row_idx]
        row.cells[col_idx].text = label
        row.cells[col_idx].paragraphs[0].runs[0].bold = True
        row.cells[col_idx+1].text = str(val or "")
        
    set_cell_text(0, 0, "姓名", record['name'])
    set_cell_text(0, 2, "工作单位", record['company'] or "暂无单位")
    set_cell_text(1, 0, "考试科目", record['exam_type'])
    
    is_pass = record['score'] >= 90
    set_cell_text(1, 2, "得分", f"{record['score']}分 ({'通过' if is_pass else '未通过'})")
    
    set_cell_text(2, 0, "答题用时", record['duration'])
    set_cell_text(2, 2, "提交时间", record['created_at'])
    
    doc.add_paragraph() # 空行
    
    # 答题详情
    h = doc.add_paragraph()
    h_run = h.add_run("答题详情记录")
    h_run.bold = True
    h_run.font.size = Pt(12)
    
    for idx, item in enumerate(details):
        p = doc.add_paragraph()
        p.add_run(f"{idx + 1}. 题目：{item['question']}\n").bold = True
        
        user_ans = item['user_answer']
        corr_ans = item['correct_answer']
        is_correct = item['is_correct'] == 1
        
        p.add_run("   您的答案：")
        user_run = p.add_run(f"{user_ans}")
        if is_correct:
            user_run.font.color.rgb = RGBColor(16, 185, 129) # 绿
        else:
            user_run.font.color.rgb = RGBColor(239, 68, 68) # 红
            
        p.add_run("    正确答案：")
        p.add_run(f"{corr_ans}").font.color.rgb = RGBColor(16, 185, 129)
        
        p.add_run("    判定结果：")
        res_run = p.add_run("回答正确" if is_correct else "回答错误")
        res_run.bold = True
        if is_correct:
            res_run.font.color.rgb = RGBColor(16, 185, 129)
        else:
            res_run.font.color.rgb = RGBColor(239, 68, 68)
            
    # 输出到 stream
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    filename = f"{record['company'] or '单位'}_{record['name']}_{record['exam_type']}答题详情.docx"
    encoded_filename = urllib.parse.quote(filename)
    headers = {
        'Content-Disposition': f"attachment; filename*=UTF-8''{encoded_filename}"
    }
    return StreamingResponse(
        file_stream, 
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
        headers=headers
    )


# ---------------- 静态页面路由 ----------------

# 托管 uploads 目录以查看照片
app.mount("/uploads", StaticFiles(directory="uploads"), name="uploads")
app.mount("/static", StaticFiles(directory="static"), name="static")

@app.get("/", response_class=HTMLResponse)
def index():
    with open("static/login.html", "r", encoding="utf-8") as f:
        return f.read()

@app.get("/login", response_class=HTMLResponse)
def get_login():
    with open("static/login.html", "r", encoding="utf-8") as f:
        return f.read()

@app.get("/client", response_class=HTMLResponse)
def get_client():
    headers = {
        "Cache-Control": "no-store, no-cache, must-revalidate, max-age=0",
        "Pragma": "no-cache",
        "Expires": "0"
    }
    with open("static/client.html", "r", encoding="utf-8") as f:
        return HTMLResponse(content=f.read(), headers=headers)

@app.get("/admin", response_class=HTMLResponse)
def get_admin():
    headers = {
        "Cache-Control": "no-store, no-cache, must-revalidate, max-age=0",
        "Pragma": "no-cache",
        "Expires": "0"
    }
    with open("static/admin.html", "r", encoding="utf-8") as f:
        return HTMLResponse(content=f.read(), headers=headers)

# 获取可用岗位/工种列表（公开接口，供客户端动态加载使用）
@app.get("/api/job_types")
def get_job_types():
    default_jobs = '普工,木工,泥工,钢筋工,电焊工,电工,安拆工,塔吊司机,吊车司机,信号工,电梯司机,管理人员,安全员'
    jobs_str = get_config('job_types', default_jobs)
    jobs_list = [j.strip() for j in jobs_str.split(',') if j.strip()]
    return {"code": 200, "job_types": jobs_list}

# 获取可用区域列表（公开接口，无需登录）
@app.get("/api/regions")
def get_regions():
    regions_str = get_config('regions', '三元肥,尿素塔')
    regions_list = [r.strip() for r in regions_str.split(',') if r.strip()]
    return {"code": 200, "regions": regions_list}

# 获取考试科目列表（无需登录，供考生和管理员使用）
@app.get("/api/exam_subjects")
def get_exam_subjects():
    subjects = get_exam_subjects_list()
    # 首次访问时把默认配置写入数据库（保持原行为）
    if not get_config('exam_subjects', ''):
        save_exam_subjects_list(subjects)
    return {"code": 200, "data": subjects}

# 增加新考试科目（仅管理员）
@app.post("/api/admin/add_exam_subject")
def add_exam_subject(name: str = Form(...), admin = Depends(get_admin_user)):
    name = name.strip()
    if not name:
        raise HTTPException(status_code=400, detail="科目名称不能为空")

    subjects = get_exam_subjects_list()

    # 检查重名
    for s in subjects:
        if s['name'] == name:
            raise HTTPException(status_code=400, detail="该科目已存在")

    # 新科目文件名直接设为 科目名.xlsx
    subjects.append({"name": name, "file": f"{name}.xlsx"})
    save_exam_subjects_list(subjects)

    return {"code": 200, "message": "科目增加成功", "data": subjects}

# 删除考试科目（仅管理员）
@app.post("/api/admin/delete_exam_subject")
def delete_exam_subject(name: str = Form(...), admin = Depends(get_admin_user)):
    name = name.strip()
    subjects = get_exam_subjects_list()

    # 过滤掉要删除的科目
    new_subjects = [s for s in subjects if s['name'] != name]
    if len(new_subjects) == len(subjects):
        raise HTTPException(status_code=404, detail="未找到该科目")

    save_exam_subjects_list(new_subjects)
    return {"code": 200, "message": "科目删除成功", "data": new_subjects}

# 获取考试配置（仅管理员）
@app.get("/api/admin/config")
def get_configs_api(admin = Depends(get_admin_user)):
    default_jobs = '普工,木工,泥工,钢筋工,电焊工,电工,安拆工,塔吊司机,吊车司机,信号工,电梯司机,管理人员,安全员'
    return {
        "code": 200,
        "data": {
            "exam_start_time": get_config('exam_start_time', '08:00:00')[:5],
            "exam_end_time": get_config('exam_end_time', '12:00:00')[:5],
            "regions": get_config('regions', '三元肥,尿素塔'),
            "job_types": get_config('job_types', default_jobs)
        }
    }

# 修改考试配置（仅管理员）
@app.post("/api/admin/config")
def save_config_api(
    start_time: str = Form(...),
    end_time: str = Form(...),
    regions: str = Form(""),
    job_types: str = Form(""),
    admin = Depends(get_admin_user)
):
    if len(start_time) == 5:
        start_time = f"{start_time}:00"
    if len(end_time) == 5:
        end_time = f"{end_time}:00"
        
    try:
        datetime.strptime(start_time, "%H:%M:%S")
        datetime.strptime(end_time, "%H:%M:%S")
    except ValueError:
        raise HTTPException(status_code=400, detail="时间格式不正确")
        
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    try:
        cursor.execute("INSERT OR REPLACE INTO configs (key, value) VALUES ('exam_start_time', ?)", (start_time,))
        cursor.execute("INSERT OR REPLACE INTO configs (key, value) VALUES ('exam_end_time', ?)", (end_time,))
        cursor.execute("INSERT OR REPLACE INTO configs (key, value) VALUES ('regions', ?)", (regions.strip(),))
        cursor.execute("INSERT OR REPLACE INTO configs (key, value) VALUES ('job_types', ?)", (job_types.strip(),))
        conn.commit()
        return {"code": 200, "message": "配置保存成功"}
    except Exception as e:
        print(f"[Error] 保存考试配置失败: {e}")
        raise HTTPException(status_code=500, detail="保存配置失败")
    finally:
        conn.close()

# 修改管理员密码（仅管理员）
@app.post("/api/admin/update_password")
def update_admin_password(
    old_password: str = Form(...),
    new_password: str = Form(...),
    admin = Depends(get_admin_user)
):
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    try:
        cursor.execute("SELECT password FROM users WHERE id = ?", (admin['id'],))
        row = cursor.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="管理员用户不存在")

        stored_pwd = row[0]
        if not verify_pwd(old_password, stored_pwd):
            raise HTTPException(status_code=400, detail="旧密码不正确")

        hashed_pwd = encrypt_pwd(new_password)
        cursor.execute("UPDATE users SET password = ? WHERE id = ?", (hashed_pwd, admin['id']))
        conn.commit()
        return {"code": 200, "message": "密码修改成功"}
    except HTTPException:
        raise  # L3: 保留具体错误，统一用 raise HTTPException 返回正确状态码
    except Exception as e:
        print(f"[Error] 修改管理员密码失败: {e}")
        raise HTTPException(status_code=500, detail="修改密码失败")
    finally:
        conn.close()

# 获取二级管理员列表（仅超级管理员）
@app.get("/api/admin/sub_admins")
def get_sub_admins(admin = Depends(get_admin_user)):
    if admin['username'] != 'admin':
        raise HTTPException(status_code=403, detail="无权执行此操作，仅限超级管理员")
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    try:
        cursor.execute("SELECT id, username, real_name, company, status, role FROM users WHERE role = 'admin' AND username != 'admin'")
        rows = cursor.fetchall()
        sub_admins = []
        for r in rows:
            sub_admins.append({
                "id": r["id"],
                "username": r["username"],
                "real_name": r["real_name"],
                "company": r["company"],
                "status": r["status"],
                "role": r["role"]
            })
        return {"code": 200, "data": sub_admins}
    except Exception as e:
        print(f"[Error] 获取二级管理员列表失败: {e}")
        raise HTTPException(status_code=500, detail="获取列表失败")
    finally:
        conn.close()

# 添加二级管理员（仅超级管理员）
@app.post("/api/admin/add_sub_admin")
def add_sub_admin(
    username: str = Form(...),
    password: str = Form(...),
    real_name: str = Form("二级管理员"),
    company: str = Form("管理部"),
    admin = Depends(get_admin_user)
):
    if admin['username'] != 'admin':
        raise HTTPException(status_code=403, detail="无权执行此操作，仅限超级管理员")
    
    username_clean = username.strip()
    password_clean = password.strip()
    if not username_clean or not password_clean:
        raise HTTPException(status_code=400, detail="用户名和密码不能为空")
    if len(password_clean) < 6:
        raise HTTPException(status_code=400, detail="密码长度不能少于6位")

    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    try:
        cursor.execute("SELECT id FROM users WHERE username = ?", (username_clean,))
        if cursor.fetchone():
            raise HTTPException(status_code=400, detail="该用户名已存在")
        
        hashed_pwd = encrypt_pwd(password_clean)
        cursor.execute(
            "INSERT INTO users (username, password, real_name, company, status, role) VALUES (?, ?, ?, ?, ?, ?)",
            (username_clean, hashed_pwd, real_name.strip(), company.strip(), 'approved', 'admin')
        )
        conn.commit()
        return {"code": 200, "message": "二级管理员添加成功"}
    except HTTPException:
        raise
    except Exception as e:
        print(f"[Error] 添加二级管理员失败: {e}")
        raise HTTPException(status_code=500, detail="添加失败")
    finally:
        conn.close()

# 删除二级管理员（仅超级管理员）
@app.post("/api/admin/delete_sub_admin")
def delete_sub_admin(
    sub_admin_id: int = Form(...),
    admin = Depends(get_admin_user)
):
    if admin['username'] != 'admin':
        raise HTTPException(status_code=403, detail="无权执行此操作，仅限超级管理员")
    
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    try:
        cursor.execute("SELECT username FROM users WHERE id = ?", (sub_admin_id,))
        row = cursor.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="该管理员不存在")
        if row[0] == 'admin':
            raise HTTPException(status_code=400, detail="不能删除超级管理员")
        
        cursor.execute("DELETE FROM users WHERE id = ?", (sub_admin_id,))
        conn.commit()
        return {"code": 200, "message": "二级管理员删除成功"}
    except HTTPException:
        raise
    except Exception as e:
        print(f"[Error] 删除二级管理员失败: {e}")
        raise HTTPException(status_code=500, detail="删除失败")
    finally:
        conn.close()

# 修改二级管理员账户名称和密码（仅超级管理员）
@app.post("/api/admin/update_sub_admin")
def update_sub_admin(
    sub_admin_id: int = Form(...),
    new_username: str = Form(...),
    new_password: str = Form(None), # 留空表示不修改
    real_name: str = Form(None),
    company: str = Form(None),
    admin = Depends(get_admin_user)
):
    if admin['username'] != 'admin':
        raise HTTPException(status_code=403, detail="无权执行此操作，仅限超级管理员")
    
    new_username_clean = new_username.strip()
    if not new_username_clean:
        raise HTTPException(status_code=400, detail="用户名不能为空")

    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    try:
        cursor.execute("SELECT username FROM users WHERE id = ?", (sub_admin_id,))
        row = cursor.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="该管理员不存在")
        if row[0] == 'admin':
            raise HTTPException(status_code=400, detail="不能在此接口修改超级管理员")

        # 检查新用户名是否冲突 (排除自身)
        cursor.execute("SELECT id FROM users WHERE username = ? AND id != ?", (new_username_clean, sub_admin_id))
        if cursor.fetchone():
            raise HTTPException(status_code=400, detail="用户名已存在")

        if new_password and new_password.strip():
            new_password_clean = new_password.strip()
            if len(new_password_clean) < 6:
                raise HTTPException(status_code=400, detail="密码长度不能少于6位")
            hashed_pwd = encrypt_pwd(new_password_clean)
            cursor.execute(
                "UPDATE users SET username = ?, password = ? WHERE id = ?",
                (new_username_clean, hashed_pwd, sub_admin_id)
            )
        else:
            cursor.execute(
                "UPDATE users SET username = ? WHERE id = ?",
                (new_username_clean, sub_admin_id)
            )
            
        if real_name is not None:
            cursor.execute("UPDATE users SET real_name = ? WHERE id = ?", (real_name.strip(), sub_admin_id))
        if company is not None:
            cursor.execute("UPDATE users SET company = ? WHERE id = ?", (company.strip(), sub_admin_id))
            
        conn.commit()
        return {"code": 200, "message": "二级管理员信息更新成功"}
    except HTTPException:
        raise
    except Exception as e:
        print(f"[Error] 更新二级管理员失败: {e}")
        raise HTTPException(status_code=500, detail="更新失败")
    finally:
        conn.close()

# 考试页面（无需登录）
@app.get("/exam", response_class=HTMLResponse)
def get_exam():
    if not is_exam_open():
        start_str = get_config('exam_start_time', '08:00:00')[:5]
        end_str = get_config('exam_end_time', '12:00:00')[:5]
        return HTMLResponse(content=f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <title>考试未开放</title>
            <style>
                body {{
                    background: #0f172a;
                    color: #e2e8f0;
                    font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif;
                    display: flex;
                    align-items: center;
                    justify-content: center;
                    height: 100vh;
                    margin: 0;
                }}
                .card {{
                    background: rgba(255, 255, 255, 0.03);
                    border: 1px solid rgba(255, 255, 255, 0.08);
                    padding: 40px;
                    border-radius: 16px;
                    text-align: center;
                    max-width: 420px;
                    box-shadow: 0 10px 30px rgba(0, 0, 0, 0.5);
                    backdrop-filter: blur(12px);
                }}
                .icon {{
                    font-size: 3.5rem;
                    margin-bottom: 20px;
                }}
                h1 {{
                    font-size: 1.6rem;
                    margin-bottom: 15px;
                    background: linear-gradient(to right, #f43f5e, #fb7185);
                    -webkit-background-clip: text;
                    -webkit-text-fill-color: transparent;
                }}
                p {{
                    font-size: 0.95rem;
                    color: #94a3b8;
                    line-height: 1.6;
                    margin: 10px 0;
                }}
                .time-box {{
                    margin-top: 25px;
                    padding: 10px 20px;
                    background: rgba(56, 189, 248, 0.08);
                    border: 1px dashed rgba(56, 189, 248, 0.3);
                    border-radius: 8px;
                    font-weight: 600;
                    color: #38bdf8;
                    display: inline-block;
                }}
            </style>
        </head>
        <body>
            <div class="card">
                <div class="icon">🚫</div>
                <h1>当前非考试开放时间</h1>
                <p>在线答题系统目前处于关闭状态。</p>
                <p>为防范乱答题和非考试时段的非授权提交，请在开放时段访问。</p>
                <div class="time-box">每天 {start_str} - {end_str}</div>
            </div>
        </body>
        </html>
        """, status_code=403)
        
    with open("static/exam.html", "r", encoding="utf-8") as f:
        return f.read()

# 诊断服务器环境的调试接口（S3: 收敛为仅管理员可访问，避免公网泄露服务器/数据库信息）
@app.get("/api/debug/check_env")
def check_env(admin = Depends(get_admin_user)):
    # 检查 python-docx 依赖
    try:
        import docx
        docx_msg = "OK"
    except Exception as e:
        docx_msg = f"Error: {str(e)}"
        
    # 检查登记卡模板文件
    template_exists = os.path.exists("登记卡.docx")
    
    # 检查 uploads 目录写入权限
    uploads_writable = False
    uploads_msg = "OK"
    try:
        os.makedirs("uploads/cards", exist_ok=True)
        test_file = "uploads/cards/test_write.txt"
        with open(test_file, "w") as f:
            f.write("test")
        os.remove(test_file)
        uploads_writable = True
    except Exception as e:
        uploads_msg = f"Error: {str(e)}"
    
    # 检查数据库文件
    db_exists = os.path.exists("peixun.db")
    
    # 检查最新 5 条记录的 Word 文件在服务器磁盘上是否真实存在
    recent_files = []
    try:
        conn = sqlite3.connect("peixun.db")
        cursor = conn.cursor()
        cursor.execute("SELECT id, name, word_path FROM records ORDER BY id DESC LIMIT 5")
        rows = cursor.fetchall()
        for r_id, r_name, r_path in rows:
            exists = False
            if r_path:
                exists = os.path.exists(r_path)
            recent_files.append({
                "id": r_id,
                "name": r_name,
                "db_word_path": r_path,
                "actually_exists_on_disk": exists
            })
        conn.close()
    except Exception as e:
        recent_files = [f"DB Query Error: {str(e)}"]
    
    return {
        "python_version": sys.version,
        "python_docx_dependency": docx_msg,
        "template_exists": template_exists,
        "uploads_writable": uploads_writable,
        "uploads_write_error": uploads_msg,
        "database_exists": db_exists,
        "latest_5_records_check": recent_files
    }

# N4: 启动入口统一由 start_server.py 负责，避免维护两份启动逻辑
# 使用方式：python start_server.py
